"""
Creating classes and building tkinter app with updater
"""

#JIAS-2020-03-0186 breaks country parsing
#JIAS-2020-03-0175 causes ROME to be added to country and MS extra data is not captured


__author__		= 'Jacob Bursavich'
__copyright__	= 'Copyright (C) 2020, Jacob Bursavich'
__credits__		= ['Jacob Bursavich']
__license__		= 'The MIT License (MIT)'
__maintainer__	= 'Jacob Bursavich'
__email__		= 'jbursavich@gmail.com'
__status__		= 'Beta'

__AppName__		= 'JIAS Automation Assistant'
__version__		= '0.2'

#LOCATION OF NEW RELEASE AND VERSION CHECK FILEs####################################################################################
location_version_check = "http://raw.githubusercontent.com/JBB-eng/TestingAutoUpdate/master/Version"
location_updated_release = "https://github.com/JBB-eng/TestingAutoUpdate/releases/download/0.2/JIAS-Automation_build0_1.exe"
####################################################################################################################################


#imports
import tkinter as tk
import pandas as pd
from tkinter import ttk, font, scrolledtext, filedialog, messagebox
from PIL import ImageTk, Image, ImageOps
from urllib.request import urlopen
from MessageBox import *
from itertools import islice
from datetime import datetime
import os, webbrowser, cgi, threading, ctypes, subprocess, time, io, re, pyperclip, time, docx
from docx import Document
from docx.shared import Pt
from ctypes import c_int, WINFUNCTYPE, windll
from ctypes.wintypes import HWND, LPCWSTR, UINT


#MS data stored in a class for easier accessiblity
class MSInfo:
	def __init__(self, authors, first_au, ms_id, title, date, ms_type, discipline, ithenticate, extra, first_co, last_co, all_co, sub_co, short_id, coi, cover_letter, files):
		self.authors = authors
		self.first_au = first_au 	
		self.ms_id = ms_id
		self.title = title 		
		self.date = date 	
		self.ms_type = ms_type
		self.discipline = discipline
		self.ithenticate = ithenticate
		self.extra = extra
		self.first_co = first_co	
		self.last_co = last_co
		self.all_co = all_co
		self.sub_co = sub_co
		self.short_id = short_id
		self.coi = coi
		self.cover_letter = []
		self.files = []



prototype = WINFUNCTYPE(c_int, HWND, LPCWSTR, LPCWSTR, UINT)
paramflags = (1, "hwnd", 0), (1, "text", "Hi"), (1, "caption", "Hello from ctypes"), (1, "flags", 0)
MessageBox = prototype(("MessageBoxW", windll.user32), paramflags)

#########
#Globals
#########


tab_names = ["New MS", "Revised MS", "Extra Tab", "blah Blah"] #add more to increase amount of tabs
tabs = [None]*len(tab_names) #holds the tab variables
download_switch = [None]*len(tab_names) #holds whether files are DLed via yes/no radio button for each tab
ms_textbox = [None]*len(tab_names) #holds the textboxes for each individual tab
ms_cover_letter = [[None] * 1 for i in range(len(tab_names))] #cover letters for each tab (just in case)
display_message = None #message that shows user processing messages, error messages, etc

folders_for_S1_check = [None]*10 #stores variables used in the S1_manuscript_check_Tool


#all countries in the world
all_countries = "Afghanistan, Albania, Algeria, Andorra, Angola, Antigua & Deps, Argentina, Armenia, Australia, Austria, Azerbaijan, Bahamas, Bahrain, Bangladesh, Barbados, Belarus, Belgium, Belize, Benin, Bhutan, Bolivia, Bosnia Herzegovina, Botswana, Brazil, Brunei, Bulgaria, Burkina, Burma, Burundi, Cambodia, Cameroon, Canada, Cape Verde, Central African Rep, Chad, Chile, China, Republic of China,Colombia, Comoros, Democratic Republic of the Congo, Republic of the Congo, Costa Rica, Côte d’Ivoire, Ivory Coast, Republic of Côte d'Ivoire, Croatia, Cuba, Cyprus, Czech Republic, Danzig, Denmark, Djibouti, Dominica, Dominican Republic, East Timor, Ecuador, Egypt, El Salvador, Equatorial Guinea, Eritrea, Estonia, Ethiopia, Fiji, Finland, France, Gabon, Gaza Strip, The Gambia, Georgia, Germany, Ghana, Greece, Grenada, Guatemala, Guinea, Guinea-Bissau, Guyana, Haiti, Holy Roman Empire, Honduras, Hungary, Iceland, India, Indonesia, Iran, Iraq, Republic of Ireland, Israel, Italy, Ivory Coast, Jamaica, Japan, Jordan, Kazakhstan, Kenya, Kiribati, North Korea, South Korea, Kosovo, Kuwait, Kyrgyzstan, Laos, Latvia, Lebanon, Lesotho, Liberia, Libya, Liechtenstein, Lithuania, Luxembourg, Macedonia, Madagascar, Malawi, Malaysia, Maldives, Mali, Malta, Marshall Islands, Mauritania, Mauritius, Mexico, Micronesia, Moldova, Monaco, Mongolia, Montenegro, Morocco, Mount Athos, Mozambique, Namibia, Nauru, Nepal, Newfoundland, Netherlands, New Zealand, Nicaragua, Niger, Nigeria, Norway, Oman, Ottoman Empire, Pakistan, Palau, Panama, Papua New Guinea,Paraguay, Peru, Philippines, Poland, Portugal, Prussia, Qatar, Romania, Russian Federation, Rwanda, St Kitts & Nevis, St Lucia, Saint Vincent & the Grenadines, Samoa, San Marino, Sao Tome & Principe, Saudi Arabia, Senegal, Serbia, Seychelles, Sierra Leone, Singapore, Slovakia, Slovenia, Solomon Islands, Somalia, South Africa, Spain, Sri Lanka, Sudan, Suriname, Swaziland, Sweden, Switzerland, Syria, Taiwan, Tajikistan, Tanzania, Thailand, Togo, Tonga, Trinidad & Tobago, Tunisia, Turkey, Turkmenistan, Tuvalu, Uganda, Ukraine, United Arab Emirates, United Kingdom, United States, Uruguay, Uzbekistan, Vanuatu, Vatican City, Venezuela, Vietnam, Yemen, Zambia, Zimbabwe".split(', ')

#multidimensional lists that hold the relevant parsed and collected data for each tab
#example: list[n][m] (n=rows, m=columns) --> list[len(tabs_names), m=?]
#all_countries.append("Tanzania")
#all_countries.append("United Kingdom")
m=20
entry1_files = [[None] * 8 for i in range(len(tab_names))] 
entry2_files = [[None] * 8 for i in range(len(tab_names))]
entry3_checkboxes = [[None] * 8 for i in range(len(tab_names))]
entry_parsed_data = [[None] * 15 for i in range(len(tab_names))]

parsing_values = [[None] * m for i in range(len(tab_names))]
parsing_bools = [[None] * m for i in range(len(tab_names))]
ms_variables_values = [[None] * m for i in range(len(tab_names))] 

#Values for parsing_values list for each individual tab
#parsing_values[0][:] = "JIAS-2020", "Submitted: ", 2, 3, 3, 3, "Submitting Author:", "Running Head:", "Author's Cover Letter:", "If you have been invited to submit an article for a supplement, please select the title of the supplement:", "Discipline:", 2, "Overall Similarity Index Percentage:"
parsing_values[0][:] = "JIAS-2020", "Submitted: ", "Title:", " (proxy) (contact)", "Wiley - Manuscript type:", "previous submission:", "Submitting Author:", "Running Head:", "Author's Cover Letter:", "If you have been invited to submit an article for a supplement, please select the title of the supplement:", "Discipline:", "Overall Similarity Index Percentage:"


def RenameFilesAndAddToFolder(shortid, method, files, firstAuthor):
	shortID = shortid
	for ndex, entry in enumerate(entry1_files[method]):
		if ndex < len(files) and entry3_checkboxes[method][ndex].get() == True:
			try:
				file_name1, file_extension1 = os.path.splitext(entry1_files[method][ndex].get())
				file_name2, file_extension2 = os.path.splitext(entry2_files[method][ndex].get())
				print("bool value " + str(ndex) + ": " + str(entry3_checkboxes[method][ndex].get()))
				os.rename(get_download_path() + '/' + entry1_files[method][ndex].get(), get_download_path() + '/' + str(firstAuthor) + " " + shortid + '/' + entry2_files[method][ndex].get() + file_extension1)
			except:
				print('not working...need to fix this!')


def check_for_files_in_dl_folder(method):
	global files, more_than_8_files
	files = []
	files.clear()
	download_path = get_download_path()
	for entry in os.listdir(download_path):
		if os.path.isfile(os.path.join(download_path, entry)):
			if not entry.startswith(('desktop.ini', 'test.xlsx', 'export.csv', 'export (', 'S1 Weekly Check')):
				files.append(entry)
				
	if len(files) > 8:
		print(files)
		print("files greater than 8")
		more_than_8_files = 1
		#print(found_file_names[method])
	else:
		print("files less than 8")
		print('files:',files)
		more_than_8_files = 0


def intersperse(lst, item):
	result = [item] * (len(lst) * 2 - 1)
	result[0::2] = lst
	return result

def get_download_path():
	"""Returns the default downloads path for linux or windows"""
	if os.name == 'nt':
		import winreg
		sub_key = r'SOFTWARE\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders'
		downloads_guid = '{374DE290-123F-4565-9164-39C4925E467B}'
		with winreg.OpenKey(winreg.HKEY_CURRENT_USER, sub_key) as key:
			location = winreg.QueryValueEx(key, downloads_guid)[0]
		return location
	else:
		return os.path.join(os.path.expanduser('~'), 'downloads')


def clear_paragraph(self, paragraph):
	p_element = paragraph._p
	p_child_elements = [elm for elm in p_element.iterchildren()]
	for child_element in p_child_elements:
		p_element.remove(child_element)


def paragraph_replace(self, search, replace, x):
	searchre = re.compile(search)
	for paragraph in self.paragraphs:
		paragraph_text = paragraph.text
		if paragraph_text:
			if searchre.search(paragraph_text):
				clear_paragraph(self, paragraph)
				para = paragraph.add_run(re.sub(search, replace, paragraph_text))
				para.font.size = Pt(10)
				paragraph.paragraph_format.space_after=Pt(0)
				if x is 2:
					para.bold = True
				else:
					para.bold = False
				paragraph.paragraph_format.line_spacing = 1.0
	return paragraph


def table_replace(self, text_value, replace):
	result = False
	tbl_regex = re.compile(text_value)
	for table in self.tables:
		for row in table.rows:
			for cell in row.cells:
				paragraphs = cell.paragraphs
				for paragraph in paragraphs:
					for run in paragraph.runs:
						font = run.font
						font.size=Pt(10)
				if cell.text:
					if tbl_regex.search(cell.text):
						cell.text = replace
						result = True
	return result




def findStringsInMiddle(a, b, text):
	return re.findall(re.escape(a)+"(.*?)"+re.escape(b),text)


#BIG PARSING FUNCTION
def bigParsingFunction (method):

	#clears all previous data from the previous data parse of the specific parsing type
	#clear_all_values[method]

	ms_temp_author_countries = [] #temporarily holds author countries 
	ms_temp_ithenticate = [] #temporarily holds the ithenticate score
	ms_temp_authors = [] #temporarily holds author names
	ms_temp_date = [] #temporarily holds submission date


	parsed = MSInfo('','','','','','','','','','','','','','','','','')

	
	if method==0:	 #New MS parsing
		print("\n@@@@@@@@\n@@@@@@@@\nRunning", tab_names[method], "function:")
		display_message.set("Successfully Parsed " + tab_names[method] + " Text!")


		"""
		Need to clear/delete the global entries here
		so that when user runs function a second time
		the old data doesn't get added to the new data

		Need to make a function with a method argument
		that clears all of the relevant data

		i.e., del entry_files1[method][:] --> del entry_files1[0][:]
		"""


		#set default names for each revelant data variable
		ms_variables_default_values = ["ms ID", "ms date", "ms title", "ms author(s)", "ms type", "ms extra data", "ms discipline", "ms ithenticate", "ms first author", "ms short ID", "ms first author country", "ms last author country", "ms all author countries"]
		#ms_variables_values = [None]*len(ms_variables_default_values) #ms_ID, short_id, ms_date, ms_title, ms_authors, ms_first_author, ms_type, ms_extra_data, ms_author_countries, ms_first_author_country, ms_last_author_country, ms_cover_letter, ms_discipline, ithenticate]

		for ndex, entry in enumerate(ms_variables_values[method]):
			if ndex < len(ms_variables_default_values):
				ms_variables_values[method][ndex] = ms_variables_default_values[ndex]

		

		#debugging: checking default names assigned to each variable
		#x=0
		#while x < len(ms_variables_values):
		#	print("ms_variables_values", x, "is", ms_variables_values[x])
		#	x += 1


		#add the text from the text box a variable
		the_MS_text = io.StringIO(ms_textbox[method].get('1.0', 'end-1c'))

		#begin parsing (if S1O changes their format, this sectoin of code may need modification)
		#parsing_codes[] should make this process much easier

		for line in the_MS_text:

			#line = line.strip()
			if parsing_values[method][0] in line or "JIAS-2019" in line:
				ms_variables_values[method][0] = line #ms id

			elif parsing_values[method][1] in line:
				ms_variables_values[method][1] = line #ms date

			elif parsing_values[method][2] in line:
				for line in islice(the_MS_text, 2):
					ms_variables_values[method][2] = line #ms title

			elif parsing_values[method][3] in line:
				ms_variables_values[method][3] = line #ms authors

			elif parsing_values[method][4] in line:
				for line in islice(the_MS_text, 2):
					ms_variables_values[method][4] = line #ms type

			elif parsing_values[method][5] in line:
				ms_variables_values[method][5] = line #ms extra data

			elif "Select Reviewers" in line:
				ms_variables_values[method][5] = line #ms extra data


			#now continuing to parse through the text

			#bool check for whether to parse for country information
			if line.startswith(parsing_values[method][6]):
				parsing_bools[method][0] = True
			elif line.startswith(parsing_values[method][7]):
				parsing_bools[method][0] = False

			#parse for country data if bool is true
			if parsing_bools[method][0]:
				for d in all_countries:
					if "Georgia" and "Atlanta," in line:
						ms_temp_author_countries.append("United States")
					elif "Georgia" and "Atlanta" in line:
						ms_temp_author_countries.append("United States")
					elif "Georgia" and "Athens," in line:
						ms_temp_author_countries.append("United States")
					elif "Georgia" and "Athens" in line:
						ms_temp_author_countries.append("United States")
					else:
						if re.search('\\b'+d+'\\b', line):
							ms_temp_author_countries.append(d) #these values will be reassigned after the parsing is completed
						#if re.search(d, line):
						#	print("gaga gaga gaga")
						#	ms_temp_author_countries.append(d) #these values will be reassigned after the parsing is completed
						

					#if (d and "Atlanta,") in line or (d and "Atlanta") in line or (d and "Athens,") in line or (d and "Athens") in line:	#this is added because of the country Georgia, which, without this condition, is added to the list when the author is from Atlanta, Georgia!
					#	ms_temp_author_countries.append("United States")
					#elif d in line:
					#	ms_temp_author_countries.append(d) #these values will be reassigned after the parsing is completed


			#bool check for whether to parse for cover letter information
			if line.startswith(parsing_values[method][8]):
				parsing_bools[method][1] = True
			elif line.startswith(parsing_values[method][9]):
				parsing_bools[method][1] = False

			#parse for cover letter data if bool is true
			if parsing_bools[method][1] == True:
				try:
					parsed.cover_letter.append(line)
					#ms_cover_letter[method].append(line)
				except:
					ms_coverletter_error = "Error: could not parse cover letter value!"
					print(ms_coverletter_error)
					display_message.set(ms_coverletter_error)

			#parsing for ms discipline value
			if re.match(parsing_values[method][10], line):
				try:
					for line in islice(the_MS_text, 2):
						ms_variables_values[method][6] = line #ms discipline
				except:
					ms_discipline_error = "Error: could not parse discipline value!"
					print(ms_discipline_error)
					display_message.set(ms_discipline_error)


			#parsing for ms ithenticate value (needs post-processing)
			if re.match(parsing_values[method][11], line):
				ms_variables_values[method][7] = line 	#ms ithenticate
				
		
		#post processing to convert ithenticate value into proper format
		if ms_variables_values[method][7] is not "ms ithenticate":
			try:
				ms_temp_ithenticate = ms_variables_values[method][7].split(':')
				ms_variables_values[method][7] = ms_temp_ithenticate[1]
				ms_variables_values[method][7] = re.sub('%', '', ms_variables_values[method][7])
				ms_variables_values[method][7] = float(ms_variables_values[method][7]) / 100
			except:
				ithenticate_error = "Error: could not perform post processing of ithenticate value!"
				print(ithenticate_error)
				display_message.set(ithenticate_error)


		#post processing of values:

		try:
			#post processing to separate first author
			ms_temp_authors = ms_variables_values[method][3].split(',')
			ms_variables_values[method][8] = ms_temp_authors[0]

			#post processing to get date value in proper format
			ms_temp_date = ms_variables_values[method][1].split(':')
			ms_temp_date = ms_temp_date[1].split(';')
			ms_variables_values[method][1] = ms_temp_date[0].strip(' ') #ms date in proper format

			#post processing to get short ID in proper format
			ms_variables_values[method][9] = re.sub('JIAS-', '', ms_variables_values[method][0]) #ms short ID in proper format

			ms_variables_values[method][13] = ms_temp_author_countries[0]	#submitting author country (is typically 1st in list, NOT 2nd)
			ms_variables_values[method][11] = ms_temp_author_countries[-1]  #last author country 

			#removes duplicates from the temporarily list of countries
			ms_temp_author_countries = list(dict.fromkeys(ms_temp_author_countries))

			#remove first line of cover letter
			#ms_cover_letter[method].pop(0)
			parsed.cover_letter.pop(0)

			print('countries:', ms_temp_author_countries)
			
			#post processing to get first and last author's country
			ms_variables_values[method][10] = ms_temp_author_countries[0] #first author country

			if len(ms_temp_author_countries) is 1:
				ms_variables_values[method][11] = ms_temp_author_countries[0]	#if only 1 author, last author country is the same as first author country
				ms_variables_values[method][13] = ms_temp_author_countries[0]	#if 1 author, submitting author country is same as 1st au
			#elif len(ms_temp_author_countries) is not 1:
				#print("countries list:", str(ms_temp_author_countries))
				#ms_variables_values[method][11] = ms_temp_author_countries[-1]	#if more than 1 author, last author country is last country in list
				#ms_variables_values[method][13] = ms_temp_author_countries[1]	#if more than 1 au, submitting au country is 2nd au in list

			ms_variables_values[method][12] = ', '.join(ms_temp_author_countries)	#all authors' countries

			#post processing of Discipline -- this section needs more details
			if re.match("Epidemiology", ms_variables_values[method][6]):
				ms_variables_values[method][6] = "BE"

		except:
			ms_variable_errors = "Error: could not perform post processing of a particular value!"
			print(ms_variable_errors)
			display_message.set(ms_variable_errors)

		try:
			temp_COI = "; " + ms_variables_values[method][3].rstrip()
			temp_COI = findStringsInMiddle(';',',', temp_COI)

			searchParameter = ''
			k=0
			while k < len(temp_COI):
				searchParameter = searchParameter + temp_COI[k] + "[AU] OR"
				k = k + 1

			if searchParameter.endswith('[AU] OR'):
				searchParameter = searchParameter[:-7]

			searchParameter = searchParameter + ("[AU]")
			searchParameter = searchParameter[1:]
			ms_variables_values[method][14] = searchParameter
		except Exception as e:
			print('could not generate COI search parameters:', e)

		try:
			parsed.ms_id = ms_variables_values[method][0].rstrip()
			parsed.date = ms_variables_values[method][1].rstrip()
			parsed.title = ms_variables_values[method][2].rstrip()
			parsed.authors = ms_variables_values[method][3].rstrip()
			parsed.ms_type = ms_variables_values[method][4].rstrip()
			parsed.extra = ms_variables_values[method][5].rstrip()
			parsed.discipline = ms_variables_values[method][6].rstrip()
			parsed.ithenticate = str(ms_variables_values[method][7]).rstrip()
			parsed.first_au = ms_variables_values[method][8].rstrip()
			parsed.short_id = ms_variables_values[method][9].rstrip()
			parsed.first_co = ms_variables_values[method][10].rstrip()
			parsed.last_co = ms_variables_values[method][11].rstrip()
			parsed.all_co = ms_variables_values[method][12].rstrip()
			parsed.sub_co = ms_variables_values[method][13].rstrip()
			parsed.coi = ms_variables_values[method][14].rstrip()



		except:
			print("error setting msInfo values")
			display_message.set("Error setting MsInfo values") 

		try:
			# Copies the appropriate data in excel format to the clipboard.  User then can CTRL+V directly
			# into the MSLOG via the online processor.  If pasting directly to Excel on their
			# laptop, then the user needs to paste into the first cell, rather than the row header
			# as pasting in the row header will make excel create a warning popup in which they
			# have to just press OK, then it will work properly
			data = parsed.authors + "	" + parsed.first_au + "	" + "	" + parsed.ms_id + "	" + parsed.title + "	" + parsed.date + "	" + parsed.ms_type + "	" + parsed.discipline + "	"  + "	" + "Editorial Assessment"  + "	"  + "	"  + "	"  + "	"  + "	"  + "	" + parsed.first_co + "	" + parsed.sub_co + "	" + parsed.last_co + "	" + parsed.all_co + "	"  + "	"  + "	"  + "	"  + "	"  + "	" + parsed.ithenticate
			pyperclip.copy(data)
			print("\nClipboard:\t\t\t", data, "\n")
		except:
			print('couldn\'t copy data properly to clipboard')

		#create the folder and add the CL to it
		try:
			if not os.path.exists(get_download_path() + '\\' + parsed.first_au + ' ' + parsed.short_id):
				os.mkdir(get_download_path() + '\\' + parsed.first_au + ' ' + parsed.short_id)
				cover_letter_to_doc = []
				#cover_letter_to_doc = intersperse(ms_cover_letter[method], '\n')
				cover_letter_to_doc = intersperse(parsed.cover_letter, '\n')
				cl_doc = Document()
				cl_doc.add_paragraph(cover_letter_to_doc)
				cl_doc.save(get_download_path() + '\\' + parsed.first_au + ' ' + parsed.short_id + '\\' + parsed.first_au + ' CL' + '.docx')
		except:
			print("failed to make MS folder and/or place CL doc inside it")


		# create the MS Details file
		try:
			regex1 = ["<<authors>>", "<<author>>", "<<id>>", 		\
			"<<title>>", "<<date>>", "<<discipline>>",	\
			'<<countries>>', '<<type>>', '<<study_design>>', \
			'<<n>>', '<<study_period>>', '<<coi_string>>']

			replace1 = [parsed.authors, parsed.first_au, parsed.short_id, parsed.title, \
					parsed.date, parsed.discipline, parsed.all_co, parsed.ms_type, \
					"study_design", "n=", "study_period", parsed.coi]

			filename = os.getcwd() + '\\Document Templates\\' + "NEW MS Details TEMPLATE.docx"

			doc = Document(filename)
			for x in range(len(regex1)):
				paragraph_replace(doc, regex1[x], replace1[x], x)
				table_replace(doc, regex1[x], replace1[x])
				
			#print('saved at:', get_download_path() + '\\' + MSInfo.first_au + ' ' + MSInfo.short_id + '\\' + MSInfo.first_au + ' MS Details.docx')
			doc.save(get_download_path() + '\\' + parsed.first_au + ' ' + parsed.short_id + '\\' + parsed.first_au + ' MS Details.docx')

		except:
			print('couldn\'t create ms details document')

		check_for_files_in_dl_folder(method)

		for x in range(len(entry1_files[method])):
			entry1_files[method][x].delete(0, 'end')
			entry2_files[method][x].delete(0, 'end')


		for ndex2, entry1 in enumerate(entry1_files[method][:]):
			if ndex2 < len(files):
				entry1.insert(0, files[ndex2])
	

		for ndex2, entry1 in enumerate(entry2_files[method][:]):
			if ndex2 < len(files):
				entry1.insert(0, parsed.first_au)

		for i in range(8):

			tk.Checkbutton(tabs[method], var=entry3_checkboxes[method][i]).grid(column=1, row=8+i, sticky='e')

		show_results_in_labels(method)


		#debugging -- check for valid data results
		print('ID:\t\t\t\t\t', parsed.ms_id)
		print('Date:\t\t\t\t', parsed.date)
		print('Title:\t\t\t\t', parsed.title)
		print('Authors:\t\t\t', parsed.authors)
		print('Type:\t\t\t\t', parsed.ms_type)
		print('Extra Data:\t\t\t', parsed.extra)
		print('Discipline:\t\t\t', parsed.discipline)
		print('Ithenticate:\t\t', parsed.ithenticate)
		print('First AU:\t\t\t', parsed.first_au)
		print('Short ID:\t\t\t', parsed.short_id)
		print('First AU Country:\t', parsed.first_co)
		print('Last AU Country:\t', parsed.last_co)
		print('All AU Countries:\t', parsed.all_co)
		print('COI parameters:\t\t', parsed.coi)




def get_download_folder():
	global download_directory
	download_directory = filedialog.askdirectory() + "/"
	folders_for_S1_check[0].configure(state="normal")
	folders_for_S1_check[0].delete(0, 'end')
	folders_for_S1_check[0].insert(0, download_directory)
	folders_for_S1_check[0].configure(state="readonly")
	s1_window.lift()

def get_editorial_folder():
	global editorial_directory
	editorial_directory = filedialog.askdirectory() + "/"
	folders_for_S1_check[1].configure(state="normal")
	folders_for_S1_check[1].delete(0, 'end')
	folders_for_S1_check[1].insert(0, editorial_directory)
	folders_for_S1_check[1].configure(state="readonly")
	s1_window.lift()

def check_for_s1_ms_in_editorial_folders():
	global missing_ms

	folders_for_S1_check[2].configure(state="normal")
	folders_for_S1_check[2].delete(0, 'end')
	folders_for_S1_check[2].insert(0, "Manuscript Check In Progress...")
	folders_for_S1_check[2].configure(state="readonly")

	files_not_found = []
	excel_exports = []
	rootdir = os.getcwd()
	editorial_dir = editorial_directory
	excel_file_dir= download_directory

	start_time = time.time()

	for subdir, dirs, files in os.walk(excel_file_dir):
		for file in files:
			#print (os.path.join(subdir, file))
			filepath = subdir + os.sep + file

			if "export" in filepath:
				excel_exports.append(filepath)

	print("Number of Excel export files found:", len(excel_exports))

	ms_IDs = [[None] * 200 for i in range(len(excel_exports))]
	clean_JIAS = lambda x : (x.replace("JIAS-", ""))

	x=0
	while x < len(excel_exports):
		if x is 0:
			excel_file = excel_file_dir + "export.csv"
			data = pd.read_csv(
				excel_file,
				#index_col=0,
				converters = {'Manuscript ID':clean_JIAS},
				usecols = ["Manuscript ID"],
				engine = "c",
				squeeze = True,

				)
			ms_IDs[x] = data.values.tolist()	

		if x is not 0:
			excel_file = excel_file_dir + "export (" + str(x) + ").csv"
			data = pd.read_csv(
				excel_file,
				#index_col=0,
				converters = {'Manuscript ID':clean_JIAS},
				usecols = ["Manuscript ID"],
				engine = "c",
				squeeze = True,
				)
			ms_IDs[x] = data.values.tolist()
		x += 1


	#multidimensional lists that hold the relevant parsed and collected data for each tab
	#example: list[n][m] (n=rows, m=columns) --> list[len(tabs_names), m=?]
	files_found = [[0] * 200 for i in range(len(ms_IDs))] 

	for x in range (len(excel_exports)):
		#for y < len(ms_IDs[x]):
		ms_IDs[x] = [sub.replace('.R1', '') for sub in ms_IDs[x]] 
		ms_IDs[x] = [sub.replace('.R2', '') for sub in ms_IDs[x]]
		ms_IDs[x] = [sub.replace('.R3', '') for sub in ms_IDs[x]] 
		

	for x in range (len(excel_exports)):
		print ("\nPost processing of export list [", str(x), "]:\n", ms_IDs[x])
		#print ("post processing of list[1]:\n", ms_IDs[1])

	for x in range (len(excel_exports)):
		for y in range (len(ms_IDs[x])):
			for subdir, dirs, files in os.walk(editorial_dir):
				for file in files:
					filepath = subdir + os.sep + file

					if str(ms_IDs[x][y]) in filepath:
						files_found[x][y] = 1
						print("[" + str(x) + "][" + str(y) + "]\t" + ms_IDs[x][y] + "\tFile Found") 
						break


	now = datetime.now()
	end_time = time.time()
	process_time = round(end_time - start_time, 2)

	time_string = now.strftime("%d-%m-%Y_%H-%M-%S")
	f= open(download_directory + "S1 Weekly Check_" + time_string + ".txt","w+")
	f.write("@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@\n")
	f.write("\tThis is the Weekly S1 Manuscript Check against the Editorial Folder\t \n")
	f.write("\t   Check performed (dd-mm-yy_hour-min-sec): " + time_string + "\t\t \n")
	f.write("\t			Time it took to process results: " + str(process_time) + " (s)\n")
	f.write("@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@\n")
	f.write("\n--DIRECTORIES SCANNED:--\n")
	f.write("-Export.csv folder:\t" + download_directory + "\n")
	f.write("-Editorial folders:\t" + editorial_directory + "\n\n")
	f.write("--LIST OF ALL CHECKED MS IDs:--\n")

	for x in range(len(excel_exports)):
		if x is 0:
			f.write("-Export.csv:\n" + str(ms_IDs[x][:]) + "\n\n")
		else:
			f.write("-Export (" + str(x) + ").csv:\n" + str(ms_IDs[x][:]) + "\n\n")

		f.write("")
	f.write("\n")
	f.write("\n\n\n--RESULTS OF S1 CHECK:--")

	for x in range(len(excel_exports)):
		if x is 0:
			f.write("\nFrom Export.csv, IDs NOT FOUND:\n")
		else:
			f.write("\nFrom Export (" + str(x) + ").csv, IDs NOT FOUND:\n")
		for y in range (len(ms_IDs[x])):
			if files_found[x][y] is 0:
				if x is 0:
					f.write("-" + ms_IDs[x][y] + "\n")
				else:
					f.write("-" + ms_IDs[x][y] + "\n")
				print("MS ID", ms_IDs[x][y], "NOT FOUND IN THE JIAS EDITORIAL FOLDERS!")
				files_not_found.append(ms_IDs[x][y])
			else:
				pass

	f.close()
	folders_for_S1_check[2].configure(state="normal")
	folders_for_S1_check[2].delete(0, 'end')
	folders_for_S1_check[2].insert(0, "MS IDs NOT FOUND: " + str(files_not_found))
	folders_for_S1_check[2].configure(state="readonly")

	#results_label_text.set("Success! Summary added to S1 Weekly Check_" + time_string + ".txt located in " + download_directory)
	#folders_for_S1_check[3].configure(text = "Success! Summary added to S1 Weekly Check_" + time_string + ".txt located in " + download_directory)


def S1_check_popup():
	"""Generate a pop-up window for special messages."""
	global s1_window
	#global results_label_text
	#results_label_text = StringVar()
	
	s1_window = tk.Tk()
	s1_window.title("Weekly S1 Manuscript Check")
	s1_window.geometry('550x125')
	rows = 0
	while rows < 50:
		s1_window.rowconfigure(rows, weight=1)
		s1_window.columnconfigure(rows, weight=1)
		rows += 1

	#s1_window.wm_attributes("-topmost", 1)

	
	download_folder_button = tk.Button(s1_window, height = "1", width = "25", text="Select Download Folder", command=get_download_folder)
	editorial_folder_button = tk.Button(s1_window, height = "1", width = "25", text="Select Editorial Folder", command=get_editorial_folder)
	run_check_button = tk.Button(s1_window, height = "2", width = "25", text="Run Weekly S1 Check", command=check_for_s1_ms_in_editorial_folders)
	

	download_folder_button.grid(column=0, row=1, columnspan = 5,)
	#"-height": must be -column, -columnspan, -in, -ipadx, -ipady, -padx, -pady, -row, -rowspan, or -sticky
	editorial_folder_button.grid(column=0, row=2, columnspan = 5,)
	run_check_button.grid(column=0, row=10, columnspan=5,)

	folders_for_S1_check[0] = tk.Entry(s1_window, width= 50)
	folders_for_S1_check[0].insert(0, "No download folder selected!")
	folders_for_S1_check[0].configure(state="readonly")
	folders_for_S1_check[0].grid(column=5, row=1, sticky='e')

	folders_for_S1_check[1] = tk.Entry(s1_window, width= 50)
	folders_for_S1_check[1].insert(0, "No editorial folder selected!")
	folders_for_S1_check[1].configure(state="readonly")
	folders_for_S1_check[1].grid(column=5, row=2, sticky='e')

	folders_for_S1_check[2] = tk.Entry(s1_window, width=50)
	folders_for_S1_check[2].insert(0, "First, Select folders of export.csv files and JIAS Editorial")
	folders_for_S1_check[2].configure(state="readonly")
	folders_for_S1_check[2].grid(column=5, row=10, sticky='e')

	folders_for_S1_check[3] = tk.Label(s1_window, text="This may take a few minutes...").grid(column=2, row=12,)
	#results_label_text.set("This may take a few minutes...")

	tk.mainloop()


#Calls parsing function
def parseText(method):
	if download_switch[method].get()==0:
		ms_file_not_downloaded_error = "Error: files NOT downloaded!"
		print(ms_file_not_downloaded_error)
		display_message.set(ms_file_not_downloaded_error)
	else:
		bigParsingFunction(method)
		#print("This will call the", tab_names[method], "function.")


def generate_copypaste_section(tab_no):
	tk.Label(tabs[tab_no], text="Add Text to Parse:", width=15, height=2).grid(column=0, row=0, sticky="w")
	ms_textbox[tab_no] =  scrolledtext.ScrolledText(tabs[tab_no], height=0, width=30)
	ms_textbox[tab_no].grid(column=0, row=1)

	#Label for DL files
	tk.Label(tabs[tab_no], text="Files downloaded?", height=1, width = 15).grid(column=0, row=2, sticky="w")
	download_switch[tab_no] = tk.IntVar()
	download_yes = tk.Radiobutton(tabs[tab_no], text="yes", value=1, variable=download_switch[tab_no])
	download_no = tk.Radiobutton(tabs[tab_no], text="no", value=0, variable=download_switch[tab_no])
	download_yes.grid(column=0, row=3, sticky="w")
	#download_yes.place(relx = 0.01, rely = 0.275)
	download_no.grid(column=0, row=3)

	#Button for parsing text
	parse_button = tk.Button(tabs[tab_no], text="Parse text", command=lambda:parseText(tab_no))
	parse_button.grid(column=0, row=3, sticky="e")


def generate_main_app_section(tab_no):
	if tab_no is 0:
		tk.Label(tabs[tab_no], text="Files to Rename:", width=15, height=1).grid(column=0, row=4, sticky="w")
		for i in range(8):
			entry1_files[tab_no][i] = tk.Entry(tabs[tab_no], width=20)
			entry1_files[tab_no][i].grid(column=0, row=8+i, sticky='w')

			entry2_files[tab_no][i] = tk.Entry(tabs[tab_no], width=20)
			entry2_files[tab_no][i].grid(column=0, row=8+i, sticky='e')

			tk.Checkbutton(tabs[tab_no], var=entry3_checkboxes[tab_no][i]).grid(column=1, row=8+i, sticky='e')
		
		lbl_list = ['ID:', 'Date:', 'Title:', 'Authors:', 'Type:', \
			'Extra:', 'Disci:', 'iThent:', '1st AU: ', \
			'ShortID:', '1AU CO:', 'LastAU CO:', 'AllAU CO:', \
			'SubmitAu CO:', 'SearchCOI:']


		for i in range(15):
			tk.Label(tabs[tab_no], text=lbl_list[i], anchor='e', width=15).grid(column=2, row=i, sticky='w')
			entry_parsed_data[tab_no][i] = tk.Entry(tabs[tab_no], width=35)
			entry_parsed_data[tab_no][i].grid(column=3, row=i, sticky='w')

		ttk.Button(tabs[tab_no], text='OK', command=RenameFiles).grid(column=1, sticky='e', row=16)	

def RenameFiles():
	pass



def show_results_in_labels(tab_no):
	for i in range(15):
		entry_parsed_data[tab_no][i].delete(0, 'end')
		entry_parsed_data[tab_no][i].insert(0, ms_variables_values[tab_no][i])



class Main:
	def __init__(self, parent):
		def CheckUpdates():
			global latest_version
			#check if __version__ is lower than latest release
			try:
				url_data = urlopen(location_version_check)
				latest_version = str(url_data.read(), 'utf-8')
				if float(__version__) < float(latest_version):
					mb = MessageBox(None,__AppName__+' '+ str(__version__)+' needs to update to version '+str(latest_version),'Update Available',flags.MB_YESNO | flags.MB_ICONQUESTION)
					if mb ==  6:
						print("picked YES")
						CallUpdateManager = UpdateManager(parent)
						pass
					elif mb == 7:
						print("Picked NO")
						pass
				else:
					#messagebox.showinfo('Software Update','No Updates are Available.')
					pass
			except Exception as e:
				messagebox.showinfo('Software Update','Unable to Check for Update, Error:' + str(e))
				#CallUpdateManager = UpdateManager(parent)

		def CheckUpdatesViaMenu():
			try:
				url_data = urlopen(location_version_check)
				latest_version = str(url_data.read(), 'utf-8')
				if float(__version__) < float(latest_version):
					mb = MessageBox(None,__AppName__+' '+ str(__version__)+' needs to update to version '+str(latest_version),'Update Available',flags.MB_YESNO | flags.MB_ICONQUESTION)
					if mb ==  6:
						print("picked YES")
						CallUpdateManager = UpdateManager(parent)
						pass
					elif mb == 7:
						print("Picked NO")
						pass
				else:
					messagebox.showinfo('Software Update','No Updates are Available.')
					pass
			except Exception as e:
				messagebox.showinfo('Software Update','Unable to Check for Update, Error:' + str(e))
				#CallUpdateManager = UpdateManager(parent)

		def AboutMe():
			#loads info
			CallDisplayAboutMe = DisplayAboutMe(parent)
			pass

		def runBinary():
			#runs an .exe file
			pass

		def UpdateUsingManager():
			#data = urllib
			#another update version
			pass

		def StartApp():

			global display_message

			CheckUpdates()
			menubar = tk.Menu(parent)
			filemenu = tk.Menu(menubar, tearoff=0)
			filemenu.add_command(label='Exit', command=parent.destroy)
			menubar.add_cascade(label='File', menu=filemenu)
			
			toolsmenu = tk.Menu(menubar, tearoff=0)
			toolsmenu.add_command(label='Weekly S1/SP MS Check', command=S1_check_popup)
			menubar.add_cascade(label='Tools', menu=toolsmenu)


			helpmenu = tk.Menu(menubar, tearoff=0)
			helpmenu.add_command(label='Check For Updates', command=CheckUpdatesViaMenu)
			helpmenu.add_command(label='About', command=AboutMe)
			menubar.add_cascade(label='Help', menu=helpmenu)
			

			parent.config(menu=menubar)

			rows = 0
			while rows < 50:
				parent.rowconfigure(rows, weight=1)
				parent.columnconfigure(rows, weight=1)
				rows += 1

			display_message = tk.StringVar() #message that shows user processing messages, error messages, etc
			display_message.set("Welcome to the JIAS Automated MS Processor!")

			#Setup of processing/error message for a more user-friendly GUI
			main_info_display = tk.Label(parent, textvariable=display_message)
			main_info_display.grid(row=500, column=25)

			#Setup for Tkinter tabs in the main window
			nb = ttk.Notebook(parent)
			nb.grid(row=1, column=1, columnspan=48, rowspan=49, sticky='NESW')

			for i in range(len(tabs)):
				tabs[i] = ttk.Frame(nb)
				rows = 0
				while rows < 50:
					tabs[i].rowconfigure(rows, weight=1)
					tabs[i].columnconfigure(rows, weight=1)
					rows += 1
				nb.add(tabs[i], text=tab_names[i])

			#Setup for the section "to copy/paste text and start parsing" for all tabs
			for x in range(len(tab_names)):
				generate_copypaste_section(x) #Adds the copy/paste text section to all tabs
				generate_main_app_section(x) #addes the custom main section for each tab


			#begins the tkinter gui application
			pass
		StartApp()


class UpdateManager(tk.Toplevel):
	def __init__(self, parent):
		tk.Toplevel.__init__(self, parent)

		self.transient(parent)
		self.result = None
		self.grab_set()
		w = 350; h = 200
		sw = self.winfo_screenwidth()
		sh = self.winfo_screenheight()
		x = (sw - w)/2
		y = (sh - h)/2
		self.geometry('{0}x{1}+{2}+{3}'.format(w, h, int(x), int(y)))
		self.resizable(width=False, height=False)
		self.title('Update Manager')
		self.wm_iconbitmap('robot.ico')

		#image = Image.open('update.png')
		#photo = ImageTk.PhotoImage(image)
		#label = tk.Label(self, image=photo)
		#label.image = photo
		#label.pack()
		#label.grid(column=0, row=0)

		def StartUpdateManager():
			#starts the download of the newer version and updates progress bar
			try:
				f=open(self.tempdir+'/'+self.appname,'wb')
				while True:
					self.newdata = self.data.read(self.chunk)
					if self.newdata:
						f.write(self.newdata)
						self.downloadeddata += self.newdata
						self.progressbar['value'] = len(self.downloadeddata)
						display_in_MBs = (self.progressbar['value'] * 0.0000001)
						self.label0.config(text=str("{:.2f}".format(self.progressbar['value'] * 0.000001)) + '/' + str("{:.2f}".format(self.filesize_text * 0.001))+' MBs')
					else:
						break
			except Exception as e:
				messagebox.showerror('Error',str(e))
				self.destroy()
			else:
				f.close()
				os.rename(self.tempdir+'/'+self.appname, __AppName__ + "v" + latest_version.rstrip() + ".exe")
				self.label0.config(text=str(str("{:.2f}".format(self.progressbar['value'] * 0.000001)) + '/' + str("{:.2f}".format(self.filesize_text * 0.001))+' MBs'))
				self.label2.config(text='Please wait a moment while application is updated...')
				self.label1.config(text='Success!')
				InstallUpdate()
							

		def InstallUpdate():
			#installs update
			#runs the downloaded newer version of the app
			#then destroy() this current version of the app

			#all future versions will also check their local working directory for early binary versions
			#of the app, and then delete them.  this will occur when the app is started.
			#also, all of the app binary files will have the following format:
			#[name of application]_v[version number].exe
			#for example: "JIASAutomationAssistant_v0.3.exe"
			OpenNewVersion = subprocess.Popen([self.tempdir+'\\'+ __AppName__+"v"+latest_version.rstrip()+".exe"])# self.appname])
			time.sleep(5)
			parent.destroy()

			pass

		#params = cgi.parse_header(self.data.headers.get('Content-Disposition', ''))
		#filename = params[-1].get('filename')
		#self.appname = filename
		#self.tempdir = os.environ.get('temp')
		#self.chunk = 1048576

		try:
			self.data = urlopen(location_updated_release)
			self.filesize = cgi.parse_header(self.data.headers.get('Content-Length', ''))[0]

			params = cgi.parse_header(self.data.headers.get('Content-Disposition', ''))
			filename = params[-1].get('filename')
			self.appname = filename
			#self.tempdir = os.environ.get('temp')
			self.tempdir = os.getcwd()
			print('temp folder:', self.tempdir)
			self.chunk = 1048576
								
		except Exception as e:
			messagebox.showerror('Error', str(e))
			self.destroy()
		else:
			self.downloadeddata = b''
			self.progressbar = ttk.Progressbar(self,
									orient='horizontal',
									length=200,
									mode='determinate',
									value=0,
									maximum=self.filesize)
			self.filesize_text = int(int(self.filesize) / 1000)
			self.label0 = ttk.Label(self, text="0 / "+str("{:.2f}".format(self.filesize_text * 0.001))+' MBs')
			self.label0.place(relx=0.5, rely=0.25, anchor=tk.CENTER)

			self.label1 = ttk.Label(self, text="Update download in progress...")
			self.label1.place(relx=0.5, rely=0.5, anchor=tk.CENTER)

			self.progressbar.place(relx=0.5, rely=0.4, anchor=tk.CENTER)

			self.label2 = ttk.Label(self, text="")
			self.label2.place(relx=0.5, rely=0.8, anchor=tk.CENTER)

			
		self.t1 = threading.Thread(target=StartUpdateManager)
		self.t1.start()	



class DisplayAboutMe(tk.Toplevel):
	def __init__(self, parent):
		tk.Toplevel.__init__(self, parent)

		self.transient(parent)
		self.result = None
		self.grab_set()
		w = 285; h = 273
		sw = self.winfo_screenwidth()
		sh = self.winfo_screenheight()
		x = (sw - w)/2
		y = (sh - h)/2
		self.geometry('{0}x{1}+{2}+{3}'.format(w, h, int(x), int(y)))
		self.resizable(width=False, height=False)
		self.title('About')
		self.wm_iconbitmap('robot.ico')

		self.image = Image.open('jias_robot1.png')
		self.size = (100, 100)
		self.thumb = ImageOps.fit(self.image, self.size, Image.ANTIALIAS)
		self.photo = ImageTk.PhotoImage(self.thumb)
		logoLabel = tk.Label(self, image=self.photo); logoLabel.pack(side=tk.TOP, pady=10)

		f1 = tk.Frame(self); f1.pack()
		f2 = tk.Frame(self); f2.pack(pady=10)
		f3 = tk.Frame(f2); f3.pack()

		def CallHyperLink(EventArgs):
			webbrowser.open_new_tab('https://ch.linkedin.com/in/jacob-bursavich')
		
		tk.Label(f1, text=__AppName__+' '+str(__version__)).pack()
		tk.Label(f1, text='Copyright (C) 2020 Jacob Bursavich').pack()
		tk.Label(f1, text='All rights reserved').pack()

		f = font.Font(size=10, slant='italic', underline=True)
		label1 = tk.Label(f3, text='jbursavich', font = f, cursor='hand2')
		label1['foreground'] = 'blue'
		label1.pack(side=tk.LEFT)
		label1.bind('<Button-1>', CallHyperLink)
		ttk.Button(self, text='OK', command=self.destroy).pack(pady=5)



def main():
	root = tk.Tk()
	root.title(__AppName__+' '+str(__version__))
	w=650; h=525
	sw = root.winfo_screenwidth()
	sh = root.winfo_screenheight()
	x = (sw - w) / 2
	y = (sh - h) / 2
	root.geometry('{0}x{1}+{2}+{3}'.format(w, h, int(x), int(y)))
	root.resizable(width=False, height=False)
	root.wm_iconbitmap('robot.ico')
	win = Main(root)
	root.mainloop()	


if __name__ == '__main__':
	main()



