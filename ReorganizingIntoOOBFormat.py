#testing git update

"""
Reorganizing updater_class_tinker_app into proper OOP style
"""

#JIAS-2020-03-0186 breaks country parsing
#JIAS-2020-03-0175 causes ROME to be added to country and MS extra data is not captured
#JIAS-2020-05-0335 TITLE IS BLANK!

__author__		= 'Jacob Bursavich'
__copyright__	= 'Copyright (C) 2020, Jacob Bursavich'
__credits__		= ['Jacob Bursavich']
__license__		= 'The MIT License (MIT)'
__maintainer__	= 'Jacob Bursavich'
__email__		= 'jbursavich@gmail.com'
__status__		= 'Beta'

__AppName__		= 'JIAS_Automation_Assistant'
__version__		= '0.3'

#LOCATION OF NEW RELEASE AND VERSION CHECK FILEs####################################################################################
location_version_check = "http://raw.githubusercontent.com/JBB-eng/TestingAutoUpdate/master/Version"
location_updated_release = "https://github.com/JBB-eng/TestingAutoUpdate/releases/download/0.2/JIAS_Automation_Assistant_v0.2.1.exe"
location_msdetails_template = "https://github.com/JBB-eng/TestingAutoUpdate/releases/download/0.2/NEW_MS_Details_TEMPLATE.docx"
location_peerreview_template = "https://github.com/JBB-eng/TestingAutoUpdate/releases/download/0.2/Author_Reviewers_Template.docx"

#if you are adding more document templates, then include the variable in this list and it will be downloaded also
document_templates_to_update = [location_peerreview_template, location_msdetails_template]
####################################################################################################################################


#imports
import tkinter as tk
import pandas as pd
import os
import webbrowser
import cgi
import threading
import ctypes
import subprocess
import time
import io
import re
import pyperclip
import docx
import requests

from tkinter import ttk, font, scrolledtext, filedialog, messagebox
from PIL import ImageTk, Image, ImageOps
from urllib.request import urlopen
from MessageBox import *
from itertools import islice
from datetime import datetime
from docx import Document
from docx.shared import Pt
from ctypes import c_int, WINFUNCTYPE, windll
from ctypes.wintypes import HWND, LPCWSTR, UINT

prototype = WINFUNCTYPE(c_int, HWND, LPCWSTR, LPCWSTR, UINT)
paramflags = (1, "hwnd", 0), (1, "text", "Hi"), (1, "caption", "Hello from ctypes"), (1, "flags", 0)
MessageBox = prototype(("MessageBoxW", windll.user32), paramflags)

#########
#Globals
#########


tab_names = ["New Manuscript Processing"] #add more to increase amount of tabs
tabs = [None]*len(tab_names) #holds the tab variables
download_switch = [None]*len(tab_names) #holds whether files are DLed via yes/no radio button for each tab
ms_textbox = [None]*len(tab_names) #holds the textboxes for each individual tab
ms_cover_letter = [[None] * 1 for i in range(len(tab_names))] #cover letters for each tab (just in case)
display_message = None #message that shows user processing messages, error messages, etc

folders_for_S1_check = [None]*10 #stores variables used in the S1_manuscript_check_Tool
folder_for_MSLogUpdate = [None]*10 #stores location of the MsLog files for the MsLog Updater Tool

#all countries in the world
all_countries = "Afghanistan, Albania, Algeria, Andorra, Angola, Antigua & Deps, Argentina, Armenia, Australia, Austria, Azerbaijan, Bahamas, Bahrain, Bangladesh, Barbados, Belarus, Belgium, Belize, Benin, Bhutan, Bolivia, Bosnia Herzegovina, Botswana, Brazil, Brunei, Bulgaria, Burkina, Burma, Burundi, Cambodia, Cameroon, Canada, Cape Verde, Central African Rep, Chad, Chile, China, Republic of China,Colombia, Comoros, Democratic Republic of the Congo, Republic of the Congo, Costa Rica, Côte d’Ivoire, Ivory Coast, Republic of Côte d'Ivoire, Croatia, Cuba, Cyprus, Czech Republic, Danzig, Denmark, Djibouti, Dominica, Dominican Republic, East Timor, Ecuador, Egypt, El Salvador, Equatorial Guinea, Eritrea, Estonia, Ethiopia, Fiji, Finland, France, Gabon, Gaza Strip, The Gambia, Georgia, Germany, Ghana, Greece, Grenada, Guatemala, Guinea, Guinea-Bissau, Guyana, Haiti, Holy Roman Empire, Honduras, Hungary, Iceland, India, Indonesia, Iran, Iraq, Republic of Ireland, Israel, Italy, Ivory Coast, Jamaica, Japan, Jordan, Kazakhstan, Kenya, Kiribati, North Korea, South Korea, Kosovo, Kuwait, Kyrgyzstan, Laos, Latvia, Lebanon, Lesotho, Liberia, Libya, Liechtenstein, Lithuania, Luxembourg, Macedonia, Madagascar, Malawi, Malaysia, Maldives, Mali, Malta, Marshall Islands, Mauritania, Mauritius, Mexico, Micronesia, Moldova, Monaco, Mongolia, Montenegro, Morocco, Mount Athos, Mozambique, Namibia, Nauru, Nepal, Newfoundland, Netherlands, New Zealand, Nicaragua, Niger, Nigeria, Norway, Oman, Ottoman Empire, Pakistan, Palau, Panama, Papua New Guinea,Paraguay, Peru, Philippines, Poland, Portugal, Prussia, Qatar, Romania, Russian Federation, Rwanda, St Kitts & Nevis, St Lucia, Saint Vincent & the Grenadines, Samoa, San Marino, Sao Tome & Principe, Saudi Arabia, Senegal, Serbia, Seychelles, Sierra Leone, Singapore, Slovakia, Slovenia, Solomon Islands, Somalia, South Africa, Spain, Sri Lanka, Sudan, Suriname, Swaziland, Sweden, Switzerland, Syria, Taiwan, Tajikistan, Tanzania, Thailand, Togo, Tonga, Trinidad & Tobago, Tunisia, Turkey, Turkmenistan, Tuvalu, Uganda, Ukraine, United Arab Emirates, United Kingdom, United States, Uruguay, Uzbekistan, Vanuatu, Vatican City, Venezuela, Vietnam, Yemen, Zambia, Zimbabwe".split(', ')

#CDC/WHO string check/country error catch (Atlanta, Georgia = USA, not Georgia (country))
cdc_who_strings = ["Centers for Disease Control", "CDC", "C.D.C", "Center for Disease Control", "Centre for Disease Control", "WHO", "W.H.O.", "World Health Organization", "World Health Organisation"]
catch_error_list = ["Atlanta", "Athens"]
global_copy_pasta = "" #global variable to hold the configured clipboard string from the MS.  I know not good practices, but easier to have as global right now

#multidimensional lists that hold the relevant parsed and collected data for each tab
#example: list[n][m] (n=rows, m=columns) --> list[len(tabs_names), m=?]
#all_countries.append("Tanzania")
#all_countries.append("United Kingdom")
m=20
entry1_files = [[None] * 8 for i in range(len(tab_names))] 
entry2_files = [[None] * 8 for i in range(len(tab_names))]
entry3_checkboxes = [[None] * 8 for i in range(len(tab_names))]
entry_parsed_data = [[None] * 15 for i in range(len(tab_names))]
parsing_values = [[None] * m for i in range(len(tab_names))]
parsing_bools = [[None] * m for i in range(len(tab_names))]


############ GLOBALS (OOP METHODS) ##############
#global_first_au = "" #assigned during PostProcessParsedData() function
files_to_ignore_in_download_folder = ["Elisa to Check Pivot Table Setups", "JAIDS_reveiw", "desktop.ini", "JIAS All manuscripts from list"] #add as many as you want here
parsing_values[0][:] = "JIAS-2020", "Submitted: ", "Title:", " (proxy) (contact)", "Wiley - Manuscript type:", "previous submission:", "Submitting Author:", "Running Head:", "Author's Cover Letter:", "If you have been invited to submit an article for a supplement, please select the title of the supplement:", "Discipline:", "Overall Similarity Index Percentage:"
#################################################

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

def intersperse(lst, item):
	result = [item] * (len(lst) * 2 - 1)
	result[0::2] = lst
	return result

def GetDownloadPath():
	"""Returns the default downloads path for linux or windows"""
	if os.name == 'nt':
		import winreg
		sub_key = r'SOFTWARE\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders'
		downloads_guid = '{374DE290-123F-4565-9164-39C4925E467B}'
		with winreg.OpenKey(winreg.HKEY_CURRENT_USER, sub_key) as key:
			location = winreg.QueryValueEx(key, downloads_guid)[0]
		return location
	else:
		return os.path.join(os.path.expanduser('~'), 'downloads')

def globalClipboardPasta():
	""" Cheap and easy way to get access to a self.variable for the tkinter button command in the UI""" 
	pyperclip.copy(global_copy_pasta)

def clear_paragraph(self, paragraph):
	p_element = paragraph._p
	p_child_elements = [elm for elm in p_element.iterchildren()]
	for child_element in p_child_elements:
		p_element.remove(child_element)

def paragraph_replace(self, search, replace, x):
	searchre = re.compile(search)
	for paragraph in self.paragraphs:
		paragraph_text = paragraph.text
		if paragraph_text:
			if searchre.search(paragraph_text):
				clear_paragraph(self, paragraph)
				para = paragraph.add_run(re.sub(search, replace, paragraph_text))
				para.font.size = Pt(10)
				paragraph.paragraph_format.space_after=Pt(0)
				if x is 2:
					para.bold = True
				else:
					para.bold = False
				paragraph.paragraph_format.line_spacing = 1.0
	return paragraph

def table_replace(self, text_value, replace):
	result = False
	tbl_regex = re.compile(text_value)
	for table in self.tables:
		for row in table.rows:
			for cell in row.cells:
				paragraphs = cell.paragraphs
				for paragraph in paragraphs:
					for run in paragraph.runs:
						font = run.font
						font.size=Pt(10)
				if cell.text:
					if tbl_regex.search(cell.text):
						cell.text = replace
						result = True
	return result

#def findStringsInMiddle(a, b, text): 
#	return re.findall(re.escape(a)+"(.*?)"+re.escape(b),text)

def findStringsInMiddle(configure, a, b, text): 
	if configure == 1:
		output = re.findall(re.escape(a)+"(.*?)"+re.escape(b),text)
	elif configure == 2:
		output = re.findall(re.escape(a)+"(.*?"+re.escape(b) + r"(?:\s*.)?)",text, flags=re.S)

	return output

#converts discpline into a shortened version (i.e., Operational and Implementation Sciences = OI)
def applyAcronymToDiscipline (discipline_phrase):
	discipline_dict = {
	'Basic and Biomedical Sciences' : 'BB',
	'Behavioural Sciences' : 'BE',
	'Clinical Sciences' : 'CS',
	'Epidemiology' : 'BE',
	'Health Economics' : 'HE',
	'Health Policy' : 'HE',
	'Humanities' : 'SH',
	'Social Sciences' : 'SH',
	'Operational and Implementation Sciences' : 'OI'
	}

	if discipline_phrase in discipline_dict:
		short_phrase = discipline_dict[discipline_phrase]
	else:
		short_phrase = discipline_phrase

	return short_phrase


#converts ms_type into a shortened version (i.e., Research Article = Res)
def applyAcronymToMsType (msType_phrase):
	mstype_dict = {
	'Research Article':'Res',
	'Short Report':'SR',
	'Review':'Rw',
	'Commentary':'Com',
	'Viewpoint':'VP',
	'Editorial':'Editorial material - Editorial',
	'Letter to the Editor':'Editorial material - Letter to editor',
	'Debate':'Editorial Material - Debate',
	'Corrigendum':'Editorial material - Corrigendum'
	}

	if msType_phrase in mstype_dict:
		short_phrase = mstype_dict[msType_phrase]
	else:
		short_phrase = msType_phrase

	return short_phrase



#MS data stored in a class for easier accessiblity
class MSInfo:
	
	#Initializer / Instance Attributes
	def __init__(self, method, authors, first_au, ms_id, title, date, ms_type, discipline, ithenticate, extra, first_co, last_co, all_co, sub_co, short_id, coi, coi2, cover_letter, parse_text, files, cdc_check):
		self.method = method
		self.authors = authors
		self.first_au = first_au 	
		self.ms_id = ms_id
		self.title = title 		
		self.date = date 	
		self.ms_type = ms_type
		self.discipline = discipline
		self.ithenticate = ithenticate
		self.extra = extra
		self.first_co = first_co	
		self.last_co = last_co
		self.all_co = all_co
		self.sub_co = sub_co
		self.short_id = short_id
		self.coi = coi #coi with author last name only
		self.coi2 = coi2 #coi with author last name and first initial (useful for Asian authors)
		self.cover_letter = []
		self.parse_text = parse_text
		self.cdc_check = cdc_check
		self.files = []

	def ParseText(self):

		if self.method is 0:	 #New MS parsing
			
			
			#set default names for each revelant data variable
			self.authors = "authors"
			self.first_au = "first_au" 	
			self.ms_id = "ms_id"
			self.title = "title" 		
			self.date = "date" 	
			self.ms_type = "ms_type"
			self.discipline = "discipline"
			self.ithenticate = "ithenticate"
			self.extra = "extra"
			self.first_co = "first_co"	
			self.last_co = "last_co"
			self.all_co = []
			self.sub_co = "sub_co"
			self.short_id = "short_id"
			self.coi = "coi"
			self.coi2 = "coi2"
			self.cover_letter = []
			self.parse_text = "parse_text"
			self.cdc_check = False
			self.files = []

			#bools for parsing
			cover_letter_bool = 0
			country_bool = 0
			msID_bool = 0
			cdc_bool = 0

			# add the text from the GUI textbox to a variable
			self.parse_text = io.StringIO(ms_textbox[self.method].get('1.0', 'end-1c'))

# BEGIN PARSING TEXT
# IF SCHOLAR ONE CHANGES THEIR FORMAT, THEN THIS SECTION
# CAN BE ADJUSTED TO FIT THE CHANGES
# THE parsing_values[] VARIABLE, SHOULD MAKE THIS PROCESS EASIER

			for line in self.parse_text:
				line = line.rstrip()

#get MS ID
				if line.startswith("ScholarOne Manuscripts"):
				 	msID_bool = 1
				elif line.startswith("Submitted:"):
					msID_bool = 0

				if msID_bool:
					if parsing_values[self.method][0] in line or "JIAS-2019" in line:
						self.ms_id = line #ms id

#if resubmission, get resubmission ID
				if "Manuscript id of previous submission: JIAS" in line:
					self.extra = line
					self.extra = self.extra[8:-5]
					

#get MS Date
				elif parsing_values[self.method][1] in line:
					self.date = line #ms date

#get MS Title
				elif parsing_values[self.method][2] in line:
					for line in islice(self.parse_text, 2):
						self.title = line.rstrip() #ms title

#get MS Authors
				elif parsing_values[self.method][3] in line:
					self.authors = line #ms authors

#get MS type
				elif parsing_values[self.method][4] in line:
					for line in islice(self.parse_text, 2):
						self.ms_type = line.rstrip() #ms type

#get MS Extra Info
#				elif parsing_values[self.method][5] in line:
#					self.extra = line #ms extra data
#
#				elif "Select Reviewers" in line:
#					self.extra = line #ms extra data
				if line.startswith("Funding Information:"): #starts at beginning of funding information (and continues into author affiliations)
					cdc_bool = 1
				elif line.startswith(parsing_values[self.method][7]): #ends at same line as end of author affiliations
					cdc_bool = 0

				if cdc_bool:
					for a in cdc_who_strings:
						#if re.search('\\b'+a+'\\b', line):
						if a in line:
							self.cdc_check = True
							#print("cdc_is true! because of the following line:", str(line))


#Get Ms Author Countries
				#bool check for whether to parse for country information
				if line.startswith(parsing_values[self.method][6]): #starts at beginning of author affiliations
					country_bool = 1
				elif line.startswith(parsing_values[self.method][7]): #ends at end of author affiliations
					country_bool = 0

				#band-aid fix for certain issues that appear when search for countries, such as "Georgia" and "Rome"
				if country_bool:
					#print(self.all_co)

					exceptions_to_avoid = ["Georgia", "Athens", "Rome"] #Georgia, the US state and other relevant stuff for the main code
					for a in exceptions_to_avoid:
						if a in line:
							print("Possible Country EXCEPTION Found Here:", line)
							break
						else:
							for b in all_countries:
								if b in line:
									self.all_co.append(b)					
					

#					for b in catch_error_list:
#						if re.search('\\b'+b+'\\b', line):
#							if re.search('Georgia', line):
#								self.all_co.append("United States")
#							else:
#								self.all_co.append("Georgia")
#
#					for c in all_countries:
#						if re.search('\\b'+c+'\\b', line):
#							self.all_co.append(c) #these values will be reassigned after the parsing is completed
#
#					for i in range(len(self.all_co) - 1):
#					    if self.all_co[i] is "Georgia" and self.all_co[i+1] is "United States":
#					        try:
#					        	self.all_co.remove(i)
#					        except Exception as a:
#					        	print("the problem:", a)

							        


#Get MS Cover Letter
				#bool check for whether to parse for cover letter information
				if line.startswith(parsing_values[self.method][8]):
					cover_letter_bool = 1
				elif line.startswith(parsing_values[self.method][9]):
					cover_letter_bool = 0

				#parse for cover letter data if bool is true
				if cover_letter_bool:
					try:
						self.cover_letter.append(line)
					except:
						cover_letter_error = "Error: could not parse cover letter value!"
						print(cover_letter_error)
						display_message.set(cover_letter_error)

#Get MS Discpline
				if re.match(parsing_values[self.method][10], line):
					try:
						for line in islice(self.parse_text, 2):
							self.discipline = line.rstrip() #ms discipline
						self.discipline = applyAcronymToDiscipline(self.discipline)
					except:
						discipline_error = "Error: could not parse discipline value!"
						print(discipline_error)
						display_message.set(discipline_error)



#Get MS Ithenticate Score (needs post-processing)
				try:
					if re.match(parsing_values[self.method][11], line):
						self.ithenticate = line 	#ms ithenticate
				except Exception as e:
					ithenticate_error = 'Could not parse ithenticate. ERROR:'
					print(ithenticate_error, e)
					display_message.set(ithenticate_error)					
					
	def PostProcessParsedData(self):

		#post processing of ithenticate
		if self.ithenticate is not "ithenticate":
			try:
				temp_ithenticate = self.ithenticate.split(':')
				temp_ithenticate = temp_ithenticate[1]
				temp_ithenticate = re.sub('%', '', temp_ithenticate)
				temp_ithenticate = float(temp_ithenticate) / 100
				self.ithenticate = temp_ithenticate
			except Exception as e:
				ithenticate_error = "Error: could not perform post processing of ithenticate value!"
				print(ithenticate_error, e)
				display_message.set(ithenticate_error)

		#post processing of first author			
		try:
			temp_authors = self.authors.split(',')
			self.first_au = temp_authors[0]
		except Exception as e:
			print("failed to post process first author. ERROR:",e)
		
		#post processing of date
		try:
			ms_temp_date = []
			ms_temp_date = self.date.split(':')
			ms_temp_date = ms_temp_date[1].split(';')
			ms_temp_date = ms_temp_date[0].strip(' ') #ms date in proper format
			self.date = ms_temp_date
		except Exception as e:
			print("failed to post process date. ERROR:", e)

		#post processing of short ID
		try:
			self.short_id = re.sub('JIAS-', '', self.ms_id) #ms short ID in proper format
			self.short_id = self.short_id
		except Exception as e:
			print("failed to post process short ID. ERROR:", e)

		#post processing of submitting author country
		try:
			self.sub_co = self.all_co[0]	#submitting author country is 1st typically, NOT 2ND)
		except Exception as e:
			print("failed to post process submitting author country. ERROR:", e)

########## THIS MAY NEED TO BE UPDATED ALONG WITH PARSING FUNCTION ###############################################

		#post processing of first author country
		try:
			self.first_co = self.all_co[0]
		except Exception as e:
			print("failed to post process first author country", e)

		##########################################################################################################

		#post processing of last author country
		try:
			self.last_co = self.all_co[-1] #last author country is last in list
		except Exception as e:
			print("failed to post process last author country", e)

		#2nd post processing of first,last,submitting author countries
		if len(self.all_co) is 1:
			try:
				self.first_co = self.all_co[0]
				self.sub_co = self.all_co[0]
				self.last_co = self.all_co[0]
			except Exception as e:
				print("failed 2nd post processing of author countries" , e)

		#post processing of all author countries
		try:
			self.all_co = list(dict.fromkeys(self.all_co)) #removes duplicates
		except Exception as e:
			print("failed to post process all author countries", e)

		#post processing of cover letter
		try:
			self.cover_letter.pop(0) #removes first entry
		except Exception as e:
			print("failed to post process cover letter", e)

########## THIS SECTION NEEDS TO BE UPDATED TO INCLUDE ALL DISCIPLINES #############################################
		
		#post processing of discipline
		if re.match("Epidemiology", self.discipline):
			self.discipline = "BE"
		############################################################################################################


	def CreateCoiString(self):
		try:
			#get coi1
			"""
			temp_coi = '; ' + self.authors
			temp_coi = findStringsInMiddle(1, ';',',', temp_coi)
			
			self.coi = ''
			k=0
			while k < len(temp_coi):
				self.coi = self.coi + temp_coi[k] + "[AU] OR"
				k = k + 1

			if self.coi.endswith('[AU] OR'):
				self.coi = self.coi[:-7]

			self.coi = self.coi + ("[AU]")
			self.coi = self.coi[1:]
			"""
			#get coi2
			temp_coi = '; ' + self.authors
			temp_coi = findStringsInMiddle(2, ';',',', temp_coi)

			self.coi = ''
			k=0
			while k < len(temp_coi):
				self.coi = self.coi + temp_coi[k] + "[AU] OR"
				k = k + 1

			if self.coi.endswith('[AU] OR'):
				self.coi = self.coi[:-7]

			self.coi = self.coi + ("[AU]")
			self.coi = self.coi[1:]
			self.coi = self.coi.replace(',' , '')


		except Exception as e:
			print('could not generate COI search parameters. ERROR TYPE:', e)

	def CopyExcelFormatToClipboard(self):
		global global_copy_pasta
		try:
			short_ms_type = applyAcronymToMsType(self.ms_type)
			data = self.authors + "	" + self.first_au + "	" + "	" + self.ms_id + "	" + self.title + "	" + self.date + "	" + short_ms_type + "	" + self.discipline + "	"  + "	" + "Editorial Assessment"  + "	"  + "	"  + "	"  + "	"  + "	"  + "	" + self.first_co + "	" + self.sub_co + "	" + self.last_co + "	" + ', '.join(self.all_co) + "	"  + "	"  + "	"  + "	"  + "	"  + "	" + str(self.ithenticate)
			global_copy_pasta = data
			pyperclip.copy(data)
		except Exception as e:
			print('failed to copy data to clipboard in excel format. ERROR:', e)


	def CreateFolderForManuscript(self):
		global ms_folder
		try:
			if not os.path.exists(GetDownloadPath() + '\\' + self.first_au + ' ' + self.short_id):
				os.mkdir(GetDownloadPath() + '\\' + self.first_au + ' ' + self.short_id)
			ms_folder = GetDownloadPath() + '\\' + self.first_au + ' ' + self.short_id + '\\'
		except Exception as e:
			print('failed to generate folder for manuscript. ERROR:', e)

	def CreateCoverLetterAndPlaceInFolder(self):
		try:
			cover_letter_to_doc = ""
			cover_letter_to_doc = intersperse(self.cover_letter, '\n')
			cover_letter_document = Document()
			cover_letter_document.add_paragraph(cover_letter_to_doc)
			cover_letter_document.save(GetDownloadPath() + '\\' + self.first_au + ' ' + self.short_id + '\\' + self.first_au + ' CL' + '.docx')
		except Exception as e:
			print('failed to create cover letter and move it to folder. ERROR:', e)

	def CreateMSDetailsAndPlaceInFolder(self):
		try:
			str_all_co = ""
			for x in range(len(self.all_co)):
				str_all_co += self.all_co[x] + ',' + ' '
			str_all_co = str_all_co[:-2]

			entries_within_doc_template = ['<<authors>>', '<<author>>', '<<id>>', 		\
			'<<title>>', '<<date>>', '<<discipline>>',	\
			'<<countries>>', '<<type>>', '<<study_design>>', \
			'<<n>>', '<<study_period>>', '<<coi_string>>', '<<ithenticate>>','<<cdc_check>>', '<<resubmission>>']

			cdc_result_string = ""
			if self.cdc_check == 1:
				cdc_result_string = "CDC or WHO affiliation found!"
			elif self.cdc_check == 0:
				cdc_result_string = "None"
			
			resubmission_string = ""
			if "JIAS-" in self.extra:
				resubmission_string = self.extra + '\r'
			else:
				pass


			replace_entries_with_this = [self.authors, self.first_au, self.short_id, self.title, \
					self.date, self.discipline, str_all_co, self.ms_type, \
					"study_design", "n=", "study_period", self.coi +'\r', str(self.ithenticate) + '\r', \
					cdc_result_string, resubmission_string]

			filename = os.getcwd() + '\\Document Templates\\' + "NEW_MS_Details_TEMPLATE.docx"

			ms_details_document = Document(filename)
			for x in range(len(entries_within_doc_template)):
				paragraph_replace(ms_details_document, entries_within_doc_template[x], replace_entries_with_this[x], x)
				table_replace(ms_details_document, entries_within_doc_template[x], replace_entries_with_this[x])
				
			ms_details_document.save(GetDownloadPath() + '\\' + self.first_au + ' ' + self.short_id + '\\' + self.first_au + ' MS Details.docx')

		except Exception as e:
			print('failed to create ms details document and move it to folder. ERROR:', e)

	def PrintParsingResults(self):
		try:
			print('ID:\t\t\t\t\t', self.ms_id)
			print('Date:\t\t\t\t', self.date)
			print('Title:\t\t\t\t', self.title)
			print('Authors:\t\t\t', self.authors)
			print('Type:\t\t\t\t', self.ms_type)
			print('Extra Data:\t\t\t', self.extra)
			print('Discipline:\t\t\t', self.discipline)
			print('Ithenticate:\t\t', self.ithenticate)
			print('First AU:\t\t\t', self.first_au)
			print('Short ID:\t\t\t', self.short_id)
			print('First AU Country:\t', self.first_co)
			print('Last AU Country:\t', self.last_co)
			print('All AU Countries:\t', self.all_co)
			print('COI parameters:\t\t', self.coi)
		except Exception as e:
			print('failed to print parsing results. ERROR:', e)

	#arguments are for files that you'd like to ignore within the download folder
	def FindFilesInDownloadFolder(self, *args):
		path = GetDownloadPath()
		for entry in os.listdir(path):
			if os.path.isfile(os.path.join(path, entry)):
				if not entry.startswith(args):
					self.files.append(entry)

		if len(self.files) >= len(entry1_files[self.method]):
			for x in range(len(entry1_files[self.method])):
				entry1_files[self.method][x].delete(0, 'end')
				entry2_files[self.method][x].delete(0, 'end')
				entry3_checkboxes[self.method][x].set(False)			
			for x in range(len(entry1_files[self.method])):
				entry3_checkboxes[self.method][x].set(True)
				entry1_files[self.method][x].insert(0, self.files[x])
				entry2_files[self.method][x].insert(0, self.first_au)
		
		if len(self.files) < len(entry1_files[self.method]):
			for x in range(len(entry1_files[self.method])):
				entry1_files[self.method][x].delete(0, 'end')
				entry2_files[self.method][x].delete(0, 'end')
				entry3_checkboxes[self.method][x].set(False)
			for x in range(len(self.files)):
				entry3_checkboxes[self.method][x].set(True)	
				entry1_files[self.method][x].insert(0, self.files[x])
				entry2_files[self.method][x].insert(0, self.first_au)


def Parser():
	global method_parsed
	parse = MSInfo(0,'','','','','','','','','','','','','','','','','','', '', 0)
	method_parsed = int(parse.method)
	parse.ParseText()
	parse.PostProcessParsedData()
	parse.CreateCoiString()

	parse.CreateFolderForManuscript()
	parse.CreateCoverLetterAndPlaceInFolder()
	parse.CreateMSDetailsAndPlaceInFolder()
	parse.FindFilesInDownloadFolder(*files_to_ignore_in_download_folder)
	parse.PrintParsingResults()
	parse.CopyExcelFormatToClipboard()
	#update gui text boxes with MS data
	text_update = [parse.ms_id, parse.date, parse.title, parse.authors, parse.ms_type, \
				parse.extra, parse.discipline, parse.ithenticate, parse.first_au, parse.short_id, \
				parse.first_co, parse.last_co, ', '.join(parse.all_co), parse.sub_co, parse.coi]

	for i in range(15):
		entry_parsed_data[0][i].configure(state="normal")
		entry_parsed_data[0][i].delete(0, 'end')
		entry_parsed_data[0][i].insert(0, text_update[i])
		entry_parsed_data[0][i].configure(state="readonly")

		#		lbl_list = ['ID:', 'Date:', 'Title:', 'Authors:', 'Type:', \
		#	'Extra:', 'Disci:', 'iThent:', '1st AU: ', \
		#	'ShortID:', '1AU CO:', 'LastAU CO:', 'AllAU CO:', \
		#	'SubmitAu CO:', 'SearchCOI:']


		#for i in range(15):
		#	tk.Label(tabs[tab_no], text=lbl_list[i], anchor='e', width=15).grid(column=2, row=i, sticky='w')
		#	entry_parsed_data[tab_no][i] = tk.Entry(tabs[tab_no], width=35)
		#	entry_parsed_data[tab_no][i].grid(column=3, row=i, sticky='w')


def RenameFilesAndAddToMsFolder():
	file_size_limit = 1000000 #this is in BYTES, i.e., 1000 = 1KB
	alert = 0
	for x in range(len(entry1_files[method_parsed])):
		if entry1_files[method_parsed][x].get() is not "" and entry3_checkboxes[method_parsed][x].get() is True:
			try:

				#if any of the files are larger than the file_size_limit, the task will continue, but a popup will notify the user
				if os.path.getsize(GetDownloadPath() + '/' + entry1_files[method_parsed][x].get()) >= file_size_limit:
					alert = 1

				file_name1, file_extension1 = os.path.splitext(entry1_files[method_parsed][x].get())
				file_name2, file_extension2 = os.path.splitext(entry2_files[method_parsed][x].get())
				os.rename(GetDownloadPath() + '/' + entry1_files[method_parsed][x].get(), ms_folder + entry2_files[method_parsed][x].get() + file_extension1)

				entry1_files[method_parsed][x].delete(0, 'end')
				entry2_files[method_parsed][x].delete(0, 'end')
				entry3_checkboxes[method_parsed][x].set(False)

			except Exception as e:
				print ('RenameFilesAndAssToMsFolder failed. ERROR:', e)
		else:
			pass

	if alert == 1:
		messagebox.showinfo('Warning: Manuscript File Sizes!','Some of the manuscript files are larger than ' + str(round(file_size_limit/1000000)) + ' MB')
			




def RenameFilesAndAddToFolder(shortid, method, files, firstAuthor):
	
	for ndex, entry in enumerate(entry1_files[method]):
		if ndex < len(files) and entry3_checkboxes[method][ndex].get() == True:
			try:
				file_name1, file_extension1 = os.path.splitext(entry1_files[method][ndex].get())
				file_name2, file_extension2 = os.path.splitext(entry2_files[method][ndex].get())
				print("bool value " + str(ndex) + ": " + str(entry3_checkboxes[method][ndex].get()))
				os.rename(GetDownloadPath() + '/' + entry1_files[method][ndex].get(), GetDownloadPath() + '/' + str(firstAuthor) + " " + shortid + '/' + entry2_files[method][ndex].get() + file_extension1)
			except:
				print('not working...need to fix this!')

def check_for_files_in_dl_folder(method):
	global files, more_than_8_files
	files = []
	files.clear()
	download_path = GetDownloadPath()
	for entry in os.listdir(download_path):
		if os.path.isfile(os.path.join(download_path, entry)):
			if not entry.startswith(('desktop.ini', 'test.xlsx', 'export.csv', 'export (', 'S1 Weekly Check')):
				files.append(entry)
				
	if len(files) > 8:
		print(files)
		print("files greater than 8")
		more_than_8_files = 1
		#print(found_file_names[method])
	else:
		print("files less than 8")
		print('files:',files)
		more_than_8_files = 0

def get_download_folder():
	global download_directory
	download_directory = filedialog.askdirectory() + "/"
	folders_for_S1_check[0].configure(state="normal")
	folders_for_S1_check[0].delete(0, 'end')
	folders_for_S1_check[0].insert(0, download_directory)
	folders_for_S1_check[0].configure(state="readonly")
	s1_window.lift()

def get_editorial_folder():
	global editorial_directory
	editorial_directory = filedialog.askdirectory() + "/"
	folders_for_S1_check[1].configure(state="normal")
	folders_for_S1_check[1].delete(0, 'end')
	folders_for_S1_check[1].insert(0, editorial_directory)
	folders_for_S1_check[1].configure(state="readonly")
	s1_window.lift()

def get_MsLog_folder():
	global MSLog_directory
	MSLog_directory = filedialog.askdirectory() + "/"
	folder_for_MSLogUpdate[0].configure(state="normal")
	folder_for_MSLogUpdate[0].delete(0, 'end')
	folder_for_MSLogUpdate[0].insert(0, MSLog_directory)
	folder_for_MSLogUpdate[0].configure(state="readonly")
	msLog_window.lift()

def slice_folder_name(my_str, sub):
	index=my_str.find(sub)
	if index != -1 :
		return my_str[index:]
	else:
		return my_str


def prepend_multiple_lines(file_name, list_of_lines):
    """Insert given list of strings as a new lines at the beginning of a file"""
 
    # define name of temporary dummy file
    dummy_file = file_name + '.bak'
    # open given original file in read mode and dummy file in write mode
    with open(file_name, 'r') as read_obj, open(dummy_file, 'w') as write_obj:
        # Iterate over the given list of strings and write them to dummy file as lines
        for line in list_of_lines:
            write_obj.write(line + '\n')
        # Read lines from original file one by one and append them to the dummy file
        for line in read_obj:
            write_obj.write(line)
 
    # remove original file
    os.remove(file_name)
    # Rename dummy file as the original file
    os.rename(dummy_file, file_name)


def delete_lines_by_condition(original_file, condition):
    """ In a file, delete the lines at line number in given list"""
 
    dummy_file = original_file + '.bak'
    is_skipped = False
    # Open original file in read only mode and dummy file in write mode
    with open(original_file, 'r') as read_obj, open(dummy_file, 'w') as write_obj:
        # Line by line copy data from original file to dummy file
        for line in read_obj:
            # if current line matches the given condition then skip that line
            if condition(line) == False:
                write_obj.write(line)
            else:
                is_skipped = True
 
    # If any line is skipped then rename dummy file as original file
    if is_skipped:
        os.remove(original_file)
        os.rename(dummy_file, original_file)
    else:
        os.remove(dummy_file)


def delete_lines_with_word(file_name, word):
    """Delete lines from a file that contains a given word / sub-string """
    delete_lines_by_condition(file_name, lambda x : word in x )

def walklevel(some_dir, level=1):
    some_dir = some_dir.rstrip(os.path.sep)
    assert os.path.isdir(some_dir)
    num_sep = some_dir.count(os.path.sep)
    for root, dirs, files in os.walk(some_dir):
        yield root, dirs, files
        num_sep_this = root.count(os.path.sep)
        if num_sep + level <= num_sep_this:
            del dirs[:]


def updateMsLog():

	#Imports
	from openpyxl import load_workbook
	from openpyxl.worksheet.datavalidation import DataValidation
	from openpyxl.utils import quote_sheetname
	
	#File names
	UpdateMSLog_File = "UpdateMSlog_2019.12.03.xlsx"
	UpdateMSLog_File = MSLog_directory + "UpdateMSlog_2019.12.03.xlsx"
	
	UpdateMSLog_File_New_Temp = "tmp.xlsx"
	UpdateMSLog_File_New_Temp = MSLog_directory + "temp.xlsx"
	
	MSLogFor2018_2020_File = "MS log for 2018-2020_fixed.xlsx"
	MSLogFor2018_2020_File = MSLog_directory + "MS log for 2018-2020_fixed.xlsx"
	
	MSLogFor2018_2020_UPDATED_FILE = "MS log for 2018-2020_updated.xlsx"
	MSLogFor2018_2020_UPDATED_FILE = MSLog_directory + "MS log for 2018-2020_updated.xlsx"

	#Load the updateLog file
	print("Loading \"%s\"" % (UpdateMSLog_File))
	wb2 = load_workbook(filename=UpdateMSLog_File, data_only=True)
	ws2 = wb2.active

	#Clean the headers in the updateLog file
	try:
		ws2.unmerge_cells('A1:G1')
	except:
		pass

	print("\t-Renaming headers")
	ws2.delete_rows(1)
	ws2["D1"].value = "Status" 
	ws2["E1"].value = "Peer review"
	ws2["F1"].value = "1st decision"
	ws2["G1"].value = "Final decision"

	#Remove all entries earlier than 2018 and any blank ones from the updateLog file
	print("\t-Removing all rows prior to 2018 (and blank entries) [takes about 30 seconds]")
	i=1
	del_rows = []
	unwanted_entries = ["2017", "2016", "2015"]
	for row in ws2.iter_rows(min_col=2, max_col=2, min_row=2):
		i += 1
		for cell in row:
			str_cell = str(cell.value)

			for x in unwanted_entries:
				if str_cell.startswith(x):
					del_rows.append(i)

			if cell.value is None:
				del_rows.append(i)

	for r in reversed(del_rows):
		ws2.delete_rows(r)

	#If "blank" Status values, then rename as Manuscript Status values in updateLog file
	print("\t-Renaming empty Status column values with Manuscript Status values")
	i=1
	row_num = []
	for row in ws2.iter_rows(min_col=4, max_col=4, min_row=2):
		for cell in row:
			i += 1
			if cell.value is None:
				row_num.append(i)

	for r in row_num:
		cell_with_value = "C" + str(r)
		cell_that_is_blank = "D" + str(r)
		ws2[cell_that_is_blank].value = ws2[cell_with_value].value


	#Remove all entries that have "Pending Payment Agreements", "Complete Checklist" or "Draft" values in the Manuscript Status column in the updateLog file
	print("\t-Removing rows with Pending Payment, etc., values in the Manuscript Status column")
	i=1
	del_rows = []
	for row in ws2.iter_rows(min_col=3, max_col=3, min_row=2):
		i += 1
		for cell in row:
			str_cell = str(cell.value)
			if str_cell.startswith("Pending Payment Agreements"):
				del_rows.append(i)

			if str_cell.startswith("Complete Checklist"):
				del_rows.append(i)

			if str_cell.startswith("Draft"):
				del_rows.append(i)

	for r in reversed(del_rows):
		ws2.delete_rows(r)


	#Rename all "status" column values to "Editorial Assessment", "Rejected", "Withdrawn", or "Accepted" in the updateLog file
	print("\t-Renaming all Status entries to: Editorial Assessment, Rejected, Withdrawn, or Accepted")
	accept = ["Accept"]
	editorial_assessment = ["Reviewer", "Invite", "Make", "Revision" ]
	withdrawn = [""]
	rejected = ["Reject"]

	for row in ws2.iter_rows(min_col=4, max_col=4, min_row=2):
		for cell in row:
			for x in rejected:
				if x in cell.value:
					cell.value = "Rejected"
			for x in editorial_assessment:
				if x in cell.value:
					cell.value = "Editorial Assessment"
			for x in accept:
				if x in cell.value:
					cell.value = "Accepted"	

	#Rename all "Peer review" column values to "Yes" or "No" in the updateLog file
	print("\t-Renaming Peer review values to: Yes or No")
	for row in ws2.iter_rows(min_col=5, max_col=5, min_row=2):
		for cell in row:
			
			try:
				cell.value = int(cell.value)
			except:
				pass
			
			if cell.value > 0:
				cell.value = "Yes"
			else:
				cell.value = "No"

	#Save the modifications made to the updateLog file as a new temporary file
	print("\t-Saving all modifications to: \"%s\"" % (UpdateMSLog_File_New_Temp))
	wb2.save(filename=UpdateMSLog_File_New_Temp)

	#Load the MSLogfor2018-2020 file
	print("\nLoading: \"%s\"" % (MSLogFor2018_2020_File))
	wb1 = load_workbook(filename=MSLogFor2018_2020_File)

	#Identify the sheets within the MSLogfor2018-2020 file
	print("\t-Sheets found in the MSLogFor2018-2020 file: " + ', '.join(wb1.sheetnames))
	ms_update_sheet = wb1["LogUpdate"]
	validation_sheet = wb1["Dropdowns and documentation"]
	ms_log_sheet = wb1["2018-2020"]
	countries_sheet = wb1["Countries"]

	#Clear all cell values from the LogUpdate sheet in the MSLogfor2018-2020 file
	print("\t-Clearing content from the LogUpdate sheet")
	for row in ms_update_sheet:
		for cell in row:
			cell.value = None

	#Load the temporary updateLog file created:
	print("\t-Loading temp file: %s" % (UpdateMSLog_File_New_Temp))
	wb2 = load_workbook(filename=UpdateMSLog_File_New_Temp)
	ws2 = wb2.active

	mr = ws2.max_row
	mc = ws2.max_column

	#Copy the modified contents of the temporary updateLog file to the LogUpdate sheet of the MSLogfor2018-2020 file
	print("\t-Copying new LogUpdate values from %s to the LogUpdate sheet of the MSLogfor2018-2020 file" % (UpdateMSLog_File_New_Temp))
	for i in range (1, mr + 1):
		for j in range (1, mc + 1):
			c = ws2.cell(row = i, column = j)

			ms_update_sheet.cell(row = i, column = j).value = c.value



	#If there are matches between the MS-IDs of the two files, then update the columns of the MSLogfor2018-2020 2018-2020 sheet to match the values within the columns of the updated LogUpdate sheet
	print("\t-Performing MATCH function and updating values for Status, Peer Review, 1st Decision, Final Decision in the MSLogfor2018-2020 file")
	list_of_MS_IDs = []
	for row in ms_update_sheet.iter_rows(min_col=1, max_col=1, min_row=2):
		for cell in row:
			list_of_MS_IDs.append(cell.value)

	row_num = 1
	for row in ms_log_sheet.iter_rows(min_col=4, max_col=4, min_row=2):
		row_num += 1
		for cell in row:
			for x in list_of_MS_IDs:
				if x == cell.value:
					y = list_of_MS_IDs.index(x) + 2
					
					stat_cell = ms_update_sheet.cell(row = y, column = 4)
					peerRev_cell = ms_update_sheet.cell(row= y, column= 5)
					firstDec_cell = ms_update_sheet.cell(row= y, column= 6)
					lastDec_cell = ms_update_sheet.cell(row= y, column= 7)
					
					ms_log_sheet.cell(row = row_num, column = 10).value = stat_cell.value
					ms_log_sheet.cell(row = row_num, column = 11).value = peerRev_cell.value
					ms_log_sheet.cell(row = row_num, column = 12).value = firstDec_cell.value
					ms_log_sheet.cell(row = row_num, column = 13).value = lastDec_cell.value


	#Update 2018-2020 Status column if there is anything in the "DOI" column so that the status value is "Published" instead of "Accepted"
	print("\t-Updating MsLog Status values to \"Published\" if a DOI value has been added")
	row_num = 1
	for row in ms_log_sheet.iter_rows(min_col = 10, max_col = 10, min_row = 2):
		row_num += 1
		for cell in row:
			doi_status = ms_log_sheet.cell(row = row_num, column = 15).value
			if doi_status is not None:
				ms_log_sheet.cell(row = row_num, column = 10).value = "Published"
			else:
				pass



	#Adding data validation to the 2018-2020 sheet of the MSLogfor2018-2020 file
	print("\t-Adding Data Validation lists to the various columns of the 2018-2020 Sheet in the MSLogfor2018-2020 file")
	dv_Status = DataValidation(type="list", formula1="='Dropdowns and documentation'!$A$3:$A$9", showDropDown = 0, allow_blank = 1)
	ms_log_sheet.add_data_validation(dv_Status)
	dv_Status.add('J2:J1048576')

	dv_ArticleCategory = DataValidation(type="list", formula1="='Dropdowns and documentation'!$B$3:$B$12", showDropDown = 0, allow_blank = 1)
	ms_log_sheet.add_data_validation(dv_ArticleCategory)
	dv_ArticleCategory.add('G2:G1048576')

	dv_Disciple = DataValidation(type="list", formula1="='Dropdowns and documentation'!$C$3:$C$9", showDropDown = 0, allow_blank = 1)
	ms_log_sheet.add_data_validation(dv_Disciple)
	dv_Disciple.add('H2:H1048576')

	dv_PartOfASupplement = DataValidation(type="list", formula1="='Dropdowns and documentation'!$D$3:$D$4", showDropDown = 0, allow_blank = 1)
	ms_log_sheet.add_data_validation(dv_PartOfASupplement)
	dv_PartOfASupplement.add('I2:I1048576')

	dv_GenderInclusion = DataValidation(type="list", formula1="='Dropdowns and documentation'!$E$3:$E$9", showDropDown = 0, allow_blank = 1)
	ms_log_sheet.add_data_validation(dv_GenderInclusion)
	dv_GenderInclusion.add('Z2:Z1048576')

	dv_CountryClassification = DataValidation(type="list", formula1="='Dropdowns and documentation'!$F$3:$F$5", showDropDown = 0, allow_blank = 1)
	ms_log_sheet.add_data_validation(dv_CountryClassification)
	dv_CountryClassification.add('W2:W1048576')

	dv_AuthorCountry = DataValidation(type="list", formula1="='Dropdowns and documentation'!$G$3:$G$4", showDropDown = 0, allow_blank = 1)
	ms_log_sheet.add_data_validation(dv_AuthorCountry)
	dv_AuthorCountry.add('T2:T1048576')
	dv_AuthorCountry.add('U2:U1048576')
	dv_AuthorCountry.add('V2:V1048576')

	dv_Countries = DataValidation(type="list", formula1="='Countries'!$A$2:$A$219", showDropDown = 0, allow_blank = 1)
	ms_log_sheet.add_data_validation(dv_Countries)
	dv_Countries.add('P2:P1048576')
	dv_Countries.add('Q2:Q1048576')
	dv_Countries.add('R2:R1048576')

	#Save the modifications made to the MSLogfor2018-2020 file to a new file
	print("\t-Saving as: \"%s\"" % (MSLogFor2018_2020_UPDATED_FILE))
	wb1.save(filename=MSLogFor2018_2020_UPDATED_FILE)

	#delete the tmp.xlsx file that was created 
	print("\t-Deleting tmp file: \"%s\"" % (UpdateMSLog_File_New_Temp))
	os.remove(UpdateMSLog_File_New_Temp)

	#print success message
	print("\n\t-Successfully updated the MSLog!")


def check_for_s1_ms_in_editorial_folders():
	global missing_ms

	folders_for_S1_check[2].configure(state="normal")
	folders_for_S1_check[2].delete(0, 'end')
	folders_for_S1_check[2].insert(0, "Manuscript Check In Progress...")
	folders_for_S1_check[2].configure(state="readonly")

	files_not_found = []
	excel_exports = []
	rootdir = os.getcwd()
	editorial_dir = editorial_directory
	excel_file_dir= download_directory
	delete_marker = "!QAZXSW@#EDC_DELETE"

	start_time = time.time()

	#for subdir, dirs, files in os.walk(excel_file_dir):
	for subdir, dirs, files in walklevel(excel_file_dir, level=0):	
		for file in files:
			#print (os.path.join(subdir, file))
			filepath = subdir + os.sep + file

			if "export" in filepath:
				excel_exports.append(filepath)

			#continue #this should keep the search within the folder selected, and exclude all of the subdirectories.

	print("Number of Excel export files found:", len(excel_exports))

	ms_IDs = [[None] * 200 for i in range(len(excel_exports))]
	ms_FirstAu = [[None] * 200 for i in range(len(excel_exports))]
	ms_FolderLocation = [[None] * 200 for i in range(len(excel_exports))]
	
	clean_JIAS = lambda x : (x.replace("JIAS-", ""))
	clean_AUTHOR = lambda x : (x.split(", ", 1))[0]

	print("excel export number of files:", str(len(excel_exports)))

	x=0
	while x < len(excel_exports):
		if x is 0:
			excel_file = excel_file_dir + "export.csv"
			data = pd.read_csv(
				excel_file,
				#index_col=0,
				converters = {'Manuscript ID':clean_JIAS},
				usecols = ["Manuscript ID"],
				engine = "c",
				squeeze = True,

				)
			data2 = pd.read_csv(
				excel_file,
				#index_col=0,
				converters = {'Submitting Author':clean_AUTHOR},
				usecols = ["Submitting Author"],
				engine = "c",
				squeeze = True,

				)
			ms_IDs[x] = data.values.tolist()
			ms_FirstAu[x] = data2.values.tolist()

		if x is not 0:
			excel_file = excel_file_dir + "export (" + str(x) + ").csv"
			data = pd.read_csv(
				excel_file,
				#index_col=0,
				converters = {'Manuscript ID':clean_JIAS},
				usecols = ["Manuscript ID"],
				engine = "c",
				squeeze = True,
				)
			data2 = pd.read_csv(
				excel_file,
				#index_col=0,
				converters = {'Submitting Author':clean_AUTHOR},
				usecols = ["Submitting Author"],
				engine = "c",
				squeeze = True,

				)
			ms_IDs[x] = data.values.tolist()
			ms_FirstAu[x] = data2.values.tolist()
		x += 1


	#multidimensional lists that hold the relevant parsed and collected data for each tab
	#example: list[n][m] (n=rows, m=columns) --> list[len(tabs_names), m=?]
	files_found = [[0] * 200 for i in range(len(ms_IDs))] 

	for x in range (len(excel_exports)):
		#for y < len(ms_IDs[x]):
		ms_IDs[x] = [sub.replace('.R1', '') for sub in ms_IDs[x]] 
		ms_IDs[x] = [sub.replace('.R2', '') for sub in ms_IDs[x]]
		ms_IDs[x] = [sub.replace('.R3', '') for sub in ms_IDs[x]] 
	
	now = datetime.now()
	time_string = now.strftime("(%d-%m-%Y) %H-%M-%S")
	the_file_name = download_directory + "S1 Weekly Check_" + time_string + ".txt"
	f= open(the_file_name,"w+", encoding="utf-8")
	f.write("-Number of ScholarOne Excel Export files found: " + str(len(excel_exports)) + "\n\n\n\n")	
	f.write("(1) MANUSCRIPTS INCLUDED IN ScholarOne EXPORT FILES:\n")


	for x in range (len(excel_exports)):
		print("\nExport file (" + str(x+1) + "):")
		f.write("\nExport file (" + str(x+1) + "):\n")
		for y in range (len(ms_IDs[x])):
			print(str(y+1) + ".\t" + ms_FirstAu[x][y], ms_IDs[x][y])
			f.write(str(y+1) + ".\t" + ms_FirstAu[x][y] + " " + ms_IDs[x][y] + "\n")
		#print ("\nPost processing of export list [", str(x), "]:\n", ms_FirstAu[x]), (ms_IDs[x])
		#print ("post processing of list[1]:\n", ms_IDs[1])

	print("\n")
	f.write('\n\n\n')

	f.write("(2) MANUSCRIPTS FROM ScholarOne EXCEL EXPORT FILES **FOUND** IN EDITORIAL FOLDERS:\n")
	for x in range (len(excel_exports)):
		for y in range (len(ms_IDs[x])):
			print("\"" + str(ms_FirstAu[x][y]) + " " + str(ms_IDs[x][y]) + "\"")
			f.write("\n\"" + str(ms_FirstAu[x][y]) + " " + str(ms_IDs[x][y]) + "\"")
			for subdir, dirs, files in os.walk(editorial_dir):
				for file in files:
					filepath = subdir + os.sep + file

					if str(ms_IDs[x][y]) in filepath:
						files_found[x][y] = 1
						
						f.write("\n")
						#print("[" + str(x) + "][" + str(y) + "]\t" + ms_IDs[x][y] + " " + ms_FirstAu[x][y] + "\tFound at\t" + filepath)
						#print(str(ms_FirstAu[x][y]) + " " + str(ms_IDs[x][y]) + "\tFound at\t" + slice_folder_name(subdir, "/Editorial/"))#filepath)  
						ms_FolderLocation[x][y] = slice_folder_name(subdir, "/Editorial")
						print("\t-Location(s): " + ms_FolderLocation[x][y])#filepath) 
						f.write("\t-Location(s): " + ms_FolderLocation[x][y])
						break
			if files_found[x][y] is not 1:
					f.write("\t" + delete_marker)
					#break
					pass

	f.write("\n\n\n\n(3) MANUSCRIPTS FROM ScholarOne EXCEL EXPORT FILES **NOT FOUND** IN EDITORIAL FOLDERS:")

	for x in range(len(excel_exports)):
		f.write("\n\nFrom Export file (" + str(x+1) + "):")
		for y in range (len(ms_IDs[x])):
			if files_found[x][y] is 0:
				f.write("\n\t" + ms_FirstAu[x][y] + " " + ms_IDs[x][y])# + "\n")
				#if x is 0:
				#	f.write("-" + ms_IDs[x][y] + "\n")
				#else:
				#	f.write("-" + ms_IDs[x][y] + "\n")
				#print("MS ID", ms_IDs[x][y], "NOT FOUND IN THE JIAS EDITORIAL FOLDERS!")
				print(ms_FirstAu[x][y] + " " + ms_IDs[x][y] + "\tNot located in the Editorial Folders.")
				files_not_found.append(ms_IDs[x][y])
			#elif files_found[x][:] is 0
				pass


	f.close()
	now = datetime.now()
	end_time = time.time()
	process_time = round(end_time - start_time, 2)


	#cleaning up date and time variable for user-friendly output to text file:
	cleanedup_time_string = ((time_string[:15] + ":" + time_string[16:])[:-3]).replace('-', '/')
	cleanedup_time_string = (cleanedup_time_string.replace('(',"")).replace(')','')
	cleanedup_time_string = cleanedup_time_string[:11] + "at " + cleanedup_time_string[11:]
	
	heading_for_txt = ["-----------------------------------------------------------", "\t\tWeekly S1 Manuscript Check", "\t  Check performed on " + cleanedup_time_string, "-----------------------------------------------------------", "-Location of ScholarOne Export files:\t" + download_directory, "-Location of JIAS Editorial folder:\t" + editorial_directory]

	prepend_multiple_lines(the_file_name, heading_for_txt)
	delete_lines_with_word(the_file_name, delete_marker)


	"""
	#time_string = now.strftime("%d-%m-%Y_%H-%M-%S")
	#f= open(download_directory + "S1 Weekly Check_" + time_string + ".txt","w+")
	f.write("@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@\n")
	f.write("\tThis is the Weekly S1 Manuscript Check against the Editorial Folder\t \n")
	f.write("\t   Check performed (dd-mm-yy_hour-min-sec): " + time_string + "\t\t \n")
	f.write("\t         Time it took to process results: " + str(process_time) + " (s)\n")
	f.write("@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@\n")
	f.write("\nFolders Scanned:--\n")
	f.write("-Directory of ScholarOne export files:\t" + download_directory + "\n")
	f.write("-Directory of JIAS Editorial folder:\t" + editorial_directory + "\n\n\n")
	f.write("List of Manuscripts contained within the ScholarOne export files:\n")

	for x in range(len(excel_exports)):
		f.write("\n\n")
		if x is 0:
			f.write("Export.csv:\n")
		elif x is not 0:
			f.write("Export (" + str(x) + ").csv:\n")
		for y in range(len(ms_IDs[x])):
			f.write("-" + str(ms_FirstAu[x][y]) + " " + str(ms_IDs[x][y] + "\n"))

	#f.write("")
	f.write("\n")
	f.write("\n\nResults of check:")

	for x in range(len(excel_exports)):
		if x is 0:
			f.write("\nFrom Export.csv,\nManuscripts NOT located in the Editorial Folders:\n")
		else:
			f.write("\nFrom Export (" + str(x) + ").csv,\nManuscripts NOT located in the Editorial Folders:\n")
		for y in range (len(ms_IDs[x])):
			if files_found[x][y] is 0:
				f.write(ms_FirstAu[x][y] + " " + ms_IDs[x][y] + "\n")
				#if x is 0:
				#	f.write("-" + ms_IDs[x][y] + "\n")
				#else:
				#	f.write("-" + ms_IDs[x][y] + "\n")
				#print("MS ID", ms_IDs[x][y], "NOT FOUND IN THE JIAS EDITORIAL FOLDERS!")
				print(ms_FirstAu[x][y] + " " + ms_IDs[x][y] + "\tNot located in the Editorial Folders.")
				files_not_found.append(ms_IDs[x][y])
			#elif files_found[x][:] is 0
				pass

	f.close()
	"""

	folders_for_S1_check[2].configure(state="normal")
	folders_for_S1_check[2].delete(0, 'end')
	folders_for_S1_check[2].insert(0, "MS IDs NOT FOUND: " + str(files_not_found))
	folders_for_S1_check[2].configure(state="readonly")

	#results_label_text.set("Success! Summary added to S1 Weekly Check_" + time_string + ".txt located in " + download_directory)
	#folders_for_S1_check[3].configure(text = "Success! Summary added to S1 Weekly Check_" + time_string + ".txt located in " + download_directory)



def MsLogUpdate_popup():
	"""Generate a pop-up window for special messages."""
	global msLog_window
	#global results_label_text
	#results_label_text = StringVar()
	
	msLog_window = tk.Tk()
	msLog_window.title("MSLog Updater")
	msLog_window.geometry('550x125')
	rows = 0
	while rows < 50:
		msLog_window.rowconfigure(rows, weight=1)
		msLog_window.columnconfigure(rows, weight=1)
		rows += 1

	#msLog_window.wm_attributes("-topmost", 1)
	
	download_folder_button = tk.Button(msLog_window, height = "1", width = "25", text="Select Folder containing Log Files", command=get_MsLog_folder)
	run_check_button = tk.Button(msLog_window, height = "2", width = "25", text="Create an Updated MSLog File", command=updateMsLog)
	

	download_folder_button.grid(column=0, row=1, columnspan = 5,)
	run_check_button.grid(column=0, row=10, columnspan=5,)

	folder_for_MSLogUpdate[0] = tk.Entry(msLog_window, width= 50)
	folder_for_MSLogUpdate[0].insert(0, "No download folder selected!")
	folder_for_MSLogUpdate[0].configure(state="readonly")
	folder_for_MSLogUpdate[0].grid(column=5, row=1, sticky='e')

	folder_for_MSLogUpdate[1] = tk.Label(msLog_window, text="This takes about 1 minutes...").grid(column=2, row=12,)
	#results_label_text.set("This may take a few minutes...")

	tk.mainloop()


def S1_check_popup():
	"""Generate a pop-up window for special messages."""
	global s1_window
	#global results_label_text
	#results_label_text = StringVar()
	
	s1_window = tk.Tk()
	s1_window.title("Weekly S1 Manuscript Check")
	s1_window.geometry('550x125')
	rows = 0
	while rows < 50:
		s1_window.rowconfigure(rows, weight=1)
		s1_window.columnconfigure(rows, weight=1)
		rows += 1

	#s1_window.wm_attributes("-topmost", 1)

	
	download_folder_button = tk.Button(s1_window, height = "1", width = "25", text="Select Download Folder", command=get_download_folder)
	editorial_folder_button = tk.Button(s1_window, height = "1", width = "25", text="Select Editorial Folder", command=get_editorial_folder)
	run_check_button = tk.Button(s1_window, height = "2", width = "25", text="Run Weekly S1 Check", command=check_for_s1_ms_in_editorial_folders)
	

	download_folder_button.grid(column=0, row=1, columnspan = 5,)
	#"-height": must be -column, -columnspan, -in, -ipadx, -ipady, -padx, -pady, -row, -rowspan, or -sticky
	editorial_folder_button.grid(column=0, row=2, columnspan = 5,)
	run_check_button.grid(column=0, row=10, columnspan=5,)

	folders_for_S1_check[0] = tk.Entry(s1_window, width= 50)
	folders_for_S1_check[0].insert(0, "No download folder selected!")
	folders_for_S1_check[0].configure(state="readonly")
	folders_for_S1_check[0].grid(column=5, row=1, sticky='e')

	folders_for_S1_check[1] = tk.Entry(s1_window, width= 50)
	folders_for_S1_check[1].insert(0, "No editorial folder selected!")
	folders_for_S1_check[1].configure(state="readonly")
	folders_for_S1_check[1].grid(column=5, row=2, sticky='e')

	folders_for_S1_check[2] = tk.Entry(s1_window, width=50)
	folders_for_S1_check[2].insert(0, "First, Select folders of export.csv files and JIAS Editorial")
	folders_for_S1_check[2].configure(state="readonly")
	folders_for_S1_check[2].grid(column=5, row=10, sticky='e')

	folders_for_S1_check[3] = tk.Label(s1_window, text="This may take a few minutes...").grid(column=2, row=12,)
	#results_label_text.set("This may take a few minutes...")

	tk.mainloop()

def parseText(method):
	if download_switch[method].get()==0:
		ms_file_not_downloaded_error = "Error: files NOT downloaded!"
		print(ms_file_not_downloaded_error)
		display_message.set(ms_file_not_downloaded_error)
	else:
		MSInfo.bigParsingFunction(method)
		#print("This will call the", tab_names[method], "function.")

def generate_copypaste_section(tab_no):

	tk.Label(tabs[tab_no], text="Paste Text from ScholarOne here:", height=1).grid(column=0, row=0, sticky="w", pady=10, padx=10)
	ms_textbox[tab_no] =  scrolledtext.ScrolledText(tabs[tab_no], height=4, width=30, wrap=tk.WORD)
	ms_textbox[tab_no].grid(column=0, row=0, rowspan=5, sticky='w', padx=10)

	#Button for parsing text
	parse_button = tk.Button(tabs[tab_no], text="(1) Process MS Info", bg='#0052cc', fg='#ffffff', anchor="w", command=Parser)
	parse_button.grid(column=0, row=4, sticky="w", ipadx= 67, padx=10)


	#Button for recopying the configured MS info to the clipboard  (I often need this!)
	clipboard_button = tk.Button(tabs[tab_no], anchor="e", text="MSLog ClipBoard", command=globalClipboardPasta)
	clipboard_button.grid(column=3, row=15, sticky="e")


def generate_main_app_section(tab_no):
	if tab_no is 0:
		tk.Label(tabs[tab_no], text="Files to Rename:", width=15, height=1).grid(column=0, row=6, sticky="w")
		for i in range(8):
			entry1_files[tab_no][i] = tk.Entry(tabs[tab_no], width=17)
			entry1_files[tab_no][i].grid(column=0, row=7+i, sticky='w', padx=10)

			entry2_files[tab_no][i] = tk.Entry(tabs[tab_no], width=22)
			entry2_files[tab_no][i].grid(column=0, row=7+i, sticky='e', ipadx=11)

			entry3_checkboxes[tab_no][i] = tk.BooleanVar()
			tk.Checkbutton(tabs[tab_no], var=entry3_checkboxes[tab_no][i]).grid(column=0, row=7+i, sticky='e')
		
		lbl_list = ['MS-ID:', 'Date:', 'Title:', 'Authors:', 'Type:', \
			'Extra:', 'Discipline:', 'iThenticate:', '1st AU: ', \
			'Short-ID:', '1st_AU CO:', 'Last_AU CO:', 'All_AUs CO:', \
			'Submitting_Au CO:', 'PUBMED_COI:']

		tk.Label(tabs[tab_no], text=lbl_list[0], anchor='e', width=20).grid(column=2, row=0, sticky='w', pady=10)
		entry_parsed_data[tab_no][0] = tk.Entry(tabs[tab_no], width=35)
		entry_parsed_data[tab_no][0].grid(column=3, row=0, sticky='w', pady=10)
		
		for i in range(1,15):
			tk.Label(tabs[tab_no], text=lbl_list[i], anchor='e', width=20).grid(column=2, row=i, sticky='w')
			entry_parsed_data[tab_no][i] = tk.Entry(tabs[tab_no], width=35)
			entry_parsed_data[tab_no][i].grid(column=3, row=i, sticky='w')

		changefiles_button = tk.Button(tabs[tab_no], text='(2) Rename Files', bg='#0052cc', fg='#ffffff', anchor="w",command=RenameFilesAndAddToMsFolder) #style='raised.TButton'
		changefiles_button.grid(column=0, row=15, sticky="w", ipadx= 74, padx=10)

def show_results_in_labels(tab_no):
	for i in range(15):
		entry_parsed_data[tab_no][i].delete(0, 'end')
		entry_parsed_data[tab_no][i].insert(0, ms_variables_values[tab_no][i])


class Main:
	def __init__(self, parent):
		def CheckUpdates():
			global latest_version
			#check if __version__ is lower than latest release
			try:
				url_data = urlopen(location_version_check)
				latest_version = str(url_data.read(), 'utf-8')
				if float(__version__) < float(latest_version):
					mb = MessageBox(None,__AppName__+' '+ str(__version__)+' needs to update to version '+str(latest_version),'Update Available',flags.MB_YESNO | flags.MB_ICONQUESTION)
					if mb ==  6:
						print("\n\nUser Chose to Update the App...\n\n")
						CallUpdateManager = UpdateManager(parent)
						pass
					elif mb == 7:
						print("\n\nUser Chose to Not Update the App...\n\n")
						pass
				else:
					#messagebox.showinfo('Software Update','No Updates are Available.')
					pass
			except Exception as e:
				messagebox.showinfo('Software Update','Unable to Check for Updates!\nTry checking your internet connection...\n\nActual Error Message:' + str(e))
				#CallUpdateManager = UpdateManager(parent)

		def CheckUpdatesViaMenu():
			try:
				url_data = urlopen(location_version_check)
				latest_version = str(url_data.read(), 'utf-8')
				if float(__version__) < float(latest_version):
					mb = MessageBox(None,__AppName__+' '+ str(__version__)+' needs to update to version '+str(latest_version),'Update Available',flags.MB_YESNO | flags.MB_ICONQUESTION)
					if mb ==  6:
						print("\n\nUser Chose to Update the App...\n\n")
						CallUpdateManager = UpdateManager(parent)
						pass
					elif mb == 7:
						print("\n\nUser Chose to Not Update the App...\n\n")
						pass
				else:
					messagebox.showinfo('Software Update','No Updates are Available.')
					pass
			except Exception as e:
				messagebox.showinfo('Software Update','Unable to Check for Update, Error:' + str(e))
				#CallUpdateManager = UpdateManager(parent)

		def AboutMe():
			#loads info
			CallDisplayAboutMe = DisplayAboutMe(parent)
			pass


		def StartApp():

			global display_message

			#check for older versions of the app in the current working directory and delete them
			#a simple 'hack' to do this, as it tries to also delete the current version but fails to do so since the file is open and so only deletes old versions.  
			f = []
			list_of_files = os.listdir(os.getcwd())

			for fname in list_of_files:
				f.append(fname)

			for file in f:
				if file.startswith(__AppName__) and file.endswith(".exe"):
					app_file_name = os.getcwd() + "\\" + file
					try:
						os.remove(app_file_name)
						print("Removed file: " + app_file_name)
					except:
						print("Did not modify the following file: " + app_file_name)
						#pass

			CheckUpdates()
			menubar = tk.Menu(parent)
			filemenu = tk.Menu(menubar, tearoff=0)
			filemenu.add_command(label='Exit', command=parent.destroy)
			menubar.add_cascade(label='File', menu=filemenu)
			
			toolsmenu = tk.Menu(menubar, tearoff=0)
			toolsmenu.add_command(label='Weekly S1/SP MS Check', command=S1_check_popup)
			toolsmenu.add_command(label='Update MSLog', command=MsLogUpdate_popup)
			menubar.add_cascade(label='Tools', menu=toolsmenu)


			helpmenu = tk.Menu(menubar, tearoff=0)
			helpmenu.add_command(label='Check For Updates', command=CheckUpdatesViaMenu)
			helpmenu.add_command(label='About', command=AboutMe)
			menubar.add_cascade(label='Help', menu=helpmenu)
			

			parent.config(menu=menubar)

			rows = 0
			while rows < 50:
				parent.rowconfigure(rows, weight=1)
				parent.columnconfigure(rows, weight=1)
				rows += 1

			display_message = tk.StringVar() #message that shows user processing messages, error messages, etc
			display_message.set("Welcome to the JIAS Automated MS Processor!")

			#Setup of processing/error message for a more user-friendly GUI
			main_info_display = tk.Label(parent, textvariable=display_message)
			main_info_display.grid(row=500, column=25)

			#Setup for Tkinter tabs in the main window
			nb = ttk.Notebook(parent)
			nb.grid(row=1, column=1, columnspan=48, rowspan=49, sticky='NESW')

			for i in range(len(tabs)):
				tabs[i] = ttk.Frame(nb)
				rows = 0
				while rows < 50:
					tabs[i].rowconfigure(rows, weight=1)
					tabs[i].columnconfigure(rows, weight=1)
					rows += 1
				nb.add(tabs[i], text=tab_names[i])

			#Setup for the section "to copy/paste text and start parsing" for all tabs
			for x in range(len(tab_names)):
				generate_copypaste_section(x) #Adds the copy/paste text section to all tabs
				generate_main_app_section(x) #addes the custom main section for each tab


			#begins the tkinter gui application
			pass
		StartApp()


class UpdateManager(tk.Toplevel):
	def __init__(self, parent):
		tk.Toplevel.__init__(self, parent)

		self.transient(parent)
		self.result = None
		self.grab_set()
		w = 350; h = 200
		sw = self.winfo_screenwidth()
		sh = self.winfo_screenheight()
		x = (sw - w)/2
		y = (sh - h)/2
		self.geometry('{0}x{1}+{2}+{3}'.format(w, h, int(x), int(y)))
		self.resizable(width=False, height=False)
		self.title('Update Manager')
		#self.wm_iconbitmap('robot.ico')
		self.wm_iconbitmap(resource_path('robot.ico'))

		#image = Image.open('update.png')
		#photo = ImageTk.PhotoImage(image)
		#label = tk.Label(self, image=photo)
		#label.image = photo
		#label.pack()
		#label.grid(column=0, row=0)

		def StartUpdateManager():
			#starts the download of the newer version and updates progress bar
			try:
				f=open(self.tempdir+'/'+self.appname,'wb')
				while True:
					self.newdata = self.data.read(self.chunk)
					if self.newdata:
						f.write(self.newdata)
						self.downloadeddata += self.newdata
						self.progressbar['value'] = len(self.downloadeddata)
						display_in_MBs = (self.progressbar['value'] * 0.0000001)
						self.label0.config(text=str("{:.2f}".format(self.progressbar['value'] * 0.000001)) + '/' + str("{:.2f}".format(self.filesize_text * 0.001))+' MBs')
					else:
						break
			except Exception as e:
				messagebox.showerror('Error',str(e))
				self.destroy()
			else:
				f.close()
				os.rename(self.tempdir+'/'+self.appname, __AppName__ + "v" + latest_version.rstrip() + ".exe")
				self.label0.config(text=str(str("{:.2f}".format(self.progressbar['value'] * 0.000001)) + '/' + str("{:.2f}".format(self.filesize_text * 0.001))+' MBs'))
				self.label2.config(text='Please wait a moment while application is updated...')
				self.label1.config(text='Success!')
				print("Finished updating main file...")
				InstallUpdate()
							

		def InstallUpdate():
			#installs update
			#runs the downloaded newer version of the app
			#then destroy() this current version of the app

			#all future versions will also check their local working directory for early binary versions
			#of the app, and then delete them.  this will occur when the app is started.
			#also, all of the app binary files will have the following format:
			#[name of application]_v[version number].exe
			#for example: "JIASAutomationAssistant_v0.3.exe"
			OpenNewVersion = subprocess.Popen([self.tempdir+'\\'+ __AppName__+"v"+latest_version.rstrip()+".exe"])# self.appname])
			time.sleep(5)
			parent.destroy()

			pass

		#params = cgi.parse_header(self.data.headers.get('Content-Disposition', ''))
		#filename = params[-1].get('filename')
		#self.appname = filename
		#self.tempdir = os.environ.get('temp')
		#self.chunk = 1048576

		try:
			print("Beginning updating for the main file...")
			self.data = urlopen(location_updated_release)
			self.filesize = cgi.parse_header(self.data.headers.get('Content-Length', ''))[0]

			params = cgi.parse_header(self.data.headers.get('Content-Disposition', ''))
			filename = params[-1].get('filename')
			self.appname = filename
			#self.tempdir = os.environ.get('temp')
			self.tempdir = os.getcwd()
			#print('temp folder:', self.tempdir)
			self.chunk = 1048576

									
		except Exception as e:
			messagebox.showerror('Error', str(e))
			self.destroy()
		else:
			self.downloadeddata = b''
			self.progressbar = ttk.Progressbar(self,
									orient='horizontal',
									length=200,
									mode='determinate',
									value=0,
									maximum=self.filesize)
			self.filesize_text = int(int(self.filesize) / 1000)
			self.label0 = ttk.Label(self, text="0 / "+str("{:.2f}".format(self.filesize_text * 0.001))+' MBs')
			self.label0.place(relx=0.5, rely=0.25, anchor=tk.CENTER)

			self.label1 = ttk.Label(self, text="Update download in progress...")
			self.label1.place(relx=0.5, rely=0.5, anchor=tk.CENTER)

			self.progressbar.place(relx=0.5, rely=0.4, anchor=tk.CENTER)

			self.label2 = ttk.Label(self, text="")
			self.label2.place(relx=0.5, rely=0.8, anchor=tk.CENTER)

			


		#also download and update the document templates
		for x in document_templates_to_update:
			try:
				url = x
				name_of_file = str(x.split("/")[-1])
				r = requests.get(url, allow_redirects=True)
				open(os.getcwd() + "\\Document Templates\\" + name_of_file, 'wb').write(r.content)
				print("Updating the document templates:", name_of_file)
			except Exception as e:
				print("There was an error:", str(e))


		self.t1 = threading.Thread(target=StartUpdateManager)
		self.t1.start()	




class DisplayAboutMe(tk.Toplevel):
	def __init__(self, parent):
		tk.Toplevel.__init__(self, parent)

		self.transient(parent)
		self.result = None
		self.grab_set()
		w = 285; h = 273
		sw = self.winfo_screenwidth()
		sh = self.winfo_screenheight()
		x = (sw - w)/2
		y = (sh - h)/2
		self.geometry('{0}x{1}+{2}+{3}'.format(w, h, int(x), int(y)))
		self.resizable(width=False, height=False)
		self.title('About')
		#self.wm_iconbitmap('robot.ico')
		self.wm_iconbitmap(resource_path('robot.ico'))

		#self.image = Image.open('jias_robot1.png')
		self.image = Image.open(resource_path('jias_robot1.png'))
		self.size = (100, 100)
		self.thumb = ImageOps.fit(self.image, self.size, Image.ANTIALIAS)
		self.photo = ImageTk.PhotoImage(self.thumb)
		logoLabel = tk.Label(self, image=self.photo); logoLabel.pack(side=tk.TOP, pady=10)

		f1 = tk.Frame(self); f1.pack()
		f2 = tk.Frame(self); f2.pack(pady=10)
		f3 = tk.Frame(f2); f3.pack()

		def CallHyperLink(EventArgs):
			webbrowser.open_new_tab('https://ch.linkedin.com/in/jacob-bursavich')
		
		tk.Label(f1, text=__AppName__+' '+str(__version__)).pack()
		tk.Label(f1, text='Copyright (C) 2020 Jacob Bursavich').pack()
		tk.Label(f1, text='All rights reserved').pack()

		f = font.Font(size=10, slant='italic', underline=True)
		label1 = tk.Label(f3, text='jbursavich', font = f, cursor='hand2')
		label1['foreground'] = 'blue'
		label1.pack(side=tk.LEFT)
		label1.bind('<Button-1>', CallHyperLink)
		ttk.Button(self, text='OK', command=self.destroy).pack(pady=5)



def main():
	root = tk.Tk()
	root.title(__AppName__+' '+str(__version__))
	w=665; h=480
	sw = root.winfo_screenwidth()
	sh = root.winfo_screenheight()
	x = (sw - w) / 2
	y = (sh - h) / 2
	root.geometry('{0}x{1}+{2}+{3}'.format(w, h, int(x), int(y)))
	root.resizable(width=False, height=False)
	#root.wm_iconbitmap('robot.ico')
	root.wm_iconbitmap(resource_path('robot.ico'))
	win = Main(root)
	root.mainloop()	


if __name__ == '__main__':
	main()



